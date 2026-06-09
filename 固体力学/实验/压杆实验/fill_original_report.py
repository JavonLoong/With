import math
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path("压杆稳定实验课堂实验报告20260518.docx")
OUT = Path("压杆稳定实验课堂实验报告20260518_原格式填写版.docx")

D = 12.0
E = 59200.0
SIGMA_ALLOW = 500.0


def calc(mu, length):
    area = math.pi * D**2 / 4
    inertia = math.pi * D**4 / 64
    section_modulus = math.pi * D**3 / 32
    radius = math.sqrt(inertia / area)
    effective_length = mu * length
    pcr = math.pi**2 * E * inertia / effective_length**2
    slenderness = effective_length / radius
    sigma_cr = pcr / area
    delta_max = section_modulus * (SIGMA_ALLOW - sigma_cr) / pcr
    return {
        "I": inertia,
        "W": section_modulus,
        "A0": area,
        "Pcr": pcr,
        "delta_max": delta_max,
        "lambda": slenderness,
        "sigma_cr": sigma_cr,
    }


cases = [
    {
        "name": "两端铰支",
        "mu": 1.0,
        "length": 891.3,
        "measured": 1194,
        "measured_note": "",
        **calc(1.0, 891.3),
    },
    {
        "name": "一端铰一端固支",
        "mu": 0.7,
        "length": 891.3,
        "measured": 1390,
        "measured_note": "（示例，按曲线替换）",
        **calc(0.7, 891.3),
    },
    {
        "name": "两端固支",
        "mu": 0.5,
        "length": 788.6,
        "measured": 3500,
        "measured_note": "（示例，按曲线替换）",
        **calc(0.5, 788.6),
    },
]


def fmt_int(value):
    return str(round(value))


def fmt_2(value):
    return f"{value:.2f}"


def error_percent(case):
    return abs(case["measured"] - case["Pcr"]) / case["Pcr"] * 100


def set_font(paragraph, size=10.5):
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)


def set_cell(cell, text, size=10):
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        set_font(paragraph, size)


def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.append(deepcopy(new_para._p))
    inserted = paragraph._parent.paragraphs[-1]
    inserted.text = text
    set_font(inserted, 10.5)
    return inserted


def paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    run = new_para.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return new_para


doc = Document(SRC)

tables = doc.tables

# 一、杆件信息
for values, table_idx in [
    (["环氧树脂", "12.00", "837.1", "59.2", "500"], 0),
]:
    for i, value in enumerate(values):
        set_cell(tables[table_idx].rows[1].cells[i], value)

# 二、公式列表
formula_values = [
    "πd^4/64",
    "πd^3/32",
    "πd^2/4",
    "π²EI/(μL)^2",
    "W([σ]-Pcr/A0)/Pcr",
    "μL/i",
    "Pcr/A0",
]
for i, value in enumerate(formula_values):
    set_cell(tables[1].rows[1].cells[i], value, 8.5)

# 三个约束状态计算表
main_table_indices = [2, 4, 6]
recheck_table_indices = [3, 5, 7]
for case, table_idx in zip(cases, main_table_indices):
    values = [
        case["name"],
        f"{case['mu']:.1f}",
        fmt_int(case["Pcr"]),
        fmt_2(case["delta_max"]),
        fmt_2(case["lambda"]),
        fmt_2(case["sigma_cr"]),
        "大柔度杆",
    ]
    for i, value in enumerate(values):
        set_cell(tables[table_idx].rows[1].cells[i], value)

for table_idx in recheck_table_indices:
    cells = tables[table_idx].rows[1].cells
    for i, value in enumerate(["未复测", "", "", "", "", "", "", ""]):
        set_cell(cells[i], value)

# 填写段落中的相当长度和实测载荷说明。
length_i = 0
pcr_i = 0
for paragraph in doc.paragraphs:
    text = paragraph.text
    if "测定压杆相当长度=" in text and length_i < len(cases):
        case = cases[length_i]
        paragraph.text = f"1）测定压杆相当长度= {case['length']:.1f} mm。"
        set_font(paragraph)
        length_i += 1
    elif "临界载荷实验测定值" in text and pcr_i < len(cases):
        case = cases[pcr_i]
        err = error_percent(case)
        paragraph.text = (
            f"3）临界载荷实验测定值Pcr= {case['measured']} N{case['measured_note']}，"
            f"误差={err:.1f}%。"
        )
        paragraph.paragraph_format.space_before = Pt(8)
        set_font(paragraph, 9.5)
        pcr_i += 1
    elif "根据测试结果和实验结果的对比" in text:
        paragraph.text = (
            "4）压杆相当长度的选择对计算结果有明显影响。Pcr与有效长度平方成反比，"
            "长度或约束系数稍有误差，计算结果就会明显变化。误差较大时应重新测量并计算。"
        )
        paragraph.paragraph_format.space_before = Pt(8)
        set_font(paragraph)

# 在第四、第五部分问题后加入可抄写答案。
answers = {
    "观察约束装置和杆件是否理想？": (
        "答：不完全理想。实际铰支或固支处存在间隙、摩擦和微小转动，杆件可能有初始弯曲，"
        "加载轴线也难以与杆件轴线完全重合，因此会使实测临界载荷与理论计算值产生偏差。"
    ),
    "观察实验曲线是否理想？": (
        "答：实验曲线不完全理想。曲线可能出现抖动、平台不明显或局部突变，主要原因是加载偏心、"
        "传感器读数波动、夹具滑移以及人工取临界点。取Pcr实应取曲线由快速上升转为近似平台的位置，"
        "不应直接取最终quit point。"
    ),
    "你认为哪个因素对实验结果影响最大": (
        "答：我认为压杆相当长度和约束条件影响最大。Pcr与有效长度的平方成反比，"
        "长度或约束系数稍有误差，计算出的临界载荷就会明显变化。"
    ),
    "谈谈自己对压杆稳定的认识": (
        "答：通过本实验认识到，细长压杆的破坏不一定是强度不足，而可能是稳定性失效。"
        "压杆临界载荷与材料弹性模量、截面惯性矩、杆长和端部约束密切相关，其中约束条件和有效长度影响很大。"
        "实验中还体会到实际装置很难达到理想边界条件，测量长度、安装对中和曲线取点都会影响结果。"
    ),
}

for paragraph in list(doc.paragraphs):
    text = paragraph.text
    for key, answer in answers.items():
        if key in text:
            paragraph_after(paragraph, answer)
            break

doc.save(OUT)
print(OUT)
