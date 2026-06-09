from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "压杆稳定实验报告_填写参考版.docx"


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_font(paragraph, size=10.5, bold=False, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 12.5, bold=True, color=(31, 78, 121))
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=(90, 90, 90))
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9)
        set_cell_shading(cell, "D9EAF7")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9)
            if widths:
                cells[i].width = Cm(widths[i])
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(1.8)
section.right_margin = Cm(1.8)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(10.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("压杆稳定实验报告填写参考版")
set_run_font(r, size=18, bold=True, color=(31, 78, 121))

add_note(
    doc,
    "说明：本文件按现场已确认数据 d=12.00 mm、环氧树脂、E=59.2 GPa、[σ]=500 MPa 计算。"
    "第二、三组 Pcr实采用合理误差示例值，可按最终实验曲线读数替换。",
)

add_heading(doc, "一、基本信息与杆件信息", 1)
add_table(
    doc,
    ["项目", "填写内容"],
    [
        ["实验日期", "2026年5月22日"],
        ["大课老师姓名-学生姓名-学号", "按原纸面实际填写"],
        ["杆件材料", "环氧树脂"],
        ["杆件直径 d", "12.00 mm"],
        ["杆件长度", "837.1 mm"],
        ["弹性模量 E", "59.2 GPa"],
        ["许可应力 [σ]", "500 MPa"],
    ],
    widths=[4, 12],
)

add_heading(doc, "二、公式与通用计算值", 1)
add_table(
    doc,
    ["项目", "公式", "本组数值"],
    [
        ["截面惯性矩 I", "I = πd^4 / 64", "1017.88 mm^4"],
        ["抗弯截面系数 W", "W = πd^3 / 32", "169.65 mm^3"],
        ["面积 A0", "A0 = πd^2 / 4", "113.10 mm^2"],
        ["临界载荷 Pcr", "Pcr = π²EI / (μL)^2", "按各约束计算"],
        ["最大允许挠度 δmax", "δmax = W([σ] - Pcr/A0) / Pcr", "按各约束计算"],
        ["柔度 λ", "λ = μL / i, i = √(I/A0)", "i = 3.00 mm"],
        ["临界应力 σcr", "σcr = Pcr/A0 = π²E/λ²", "按各约束计算"],
    ],
    widths=[3.2, 7.8, 4.2],
)
add_note(doc, "注意：曲线上取点时的横向位移不是表格中的 δmax。δmax 是按许可应力计算出的理论最大允许挠度。")

add_heading(doc, "三、实验数据处理与计算分析", 1)
add_table(
    doc,
    ["约束状态", "μ", "压杆相当长度 L/mm", "Pcr计/N", "δmax/mm", "λ", "σcr/MPa", "压杆种类"],
    [
        ["两端铰支", "1.0", "891.3", "749", "111.80", "297.10", "6.62", "大柔度杆"],
        ["一端铰一端固支", "0.7", "891.3", "1528", "54.02", "207.97", "13.51", "大柔度杆"],
        ["两端固支", "0.5", "788.6", "3825", "20.67", "131.43", "33.82", "大柔度杆"],
    ],
    widths=[3.2, 1.1, 2.8, 2.0, 2.0, 1.6, 1.9, 2.0],
)
add_table(
    doc,
    ["约束状态", "Pcr计/N", "Pcr实/N", "误差", "备注"],
    [
        ["两端铰支", "749", "1194", "59.5%", "按前面曲线平台开始附近取点"],
        ["一端铰一端固支", "1528", "1390（示例）", "9.0%", "有实际曲线时替换"],
        ["两端固支", "3825", "3500（示例）", "8.5%", "有实际曲线时替换"],
    ],
    widths=[3.2, 2.0, 2.5, 1.7, 6.0],
)
add_note(
    doc,
    "第一组 Pcr实=1194 N 来自前面曲线平台开始附近取点。第二、三组如已有实际曲线，应把示例值替换成曲线拐点或平台开始处对应载荷。",
)

add_heading(doc, "四、各组可抄写内容", 1)
for title, lines in [
    (
        "1. 两端铰支",
        [
            "测定压杆相当长度 = 891.3 mm。",
            "该状态相关计算：两端铰支，μ=1.0，Pcr计=749 N，δmax=111.80 mm，λ=297.10，σcr=6.62 MPa，压杆种类为大柔度杆。",
            "临界载荷实验测定值 Pcr实=1194 N，误差 = |1194-749|/749×100% = 59.5%。",
            "该组误差较大，说明实际约束、相当长度测量或临界点取值对计算结果影响明显，应结合曲线重新核对取点。",
        ],
    ),
    (
        "2. 一端铰一端固支",
        [
            "测定压杆相当长度 = 891.3 mm。",
            "该状态相关计算：一端铰一端固支，μ=0.7，Pcr计=1528 N，δmax=54.02 mm，λ=207.97，σcr=13.51 MPa，压杆种类为大柔度杆。",
            "若按合理误差示例取 Pcr实=1390 N，则误差 = |1390-1528|/1528×100% = 9.0%。",
            "如最终曲线读数不同，应替换 Pcr实 并重新计算误差。",
        ],
    ),
    (
        "3. 两端固支",
        [
            "测定压杆相当长度 = 788.6 mm。",
            "该状态相关计算：两端固支，μ=0.5，Pcr计=3825 N，δmax=20.67 mm，λ=131.43，σcr=33.82 MPa，压杆种类为大柔度杆。",
            "若按合理误差示例取 Pcr实=3500 N，则误差 = |3500-3825|/3825×100% = 8.5%。",
            "如最终曲线读数不同，应替换 Pcr实 并重新计算误差。",
        ],
    ),
]:
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True)
    for line in lines:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.6)
        r = p.add_run(line)
        set_run_font(r, size=10.5)

add_heading(doc, "五、观察实验曲线与结果分析", 1)
answers = [
    (
        "1）观察约束装置和杆件是否理想？",
        "约束装置和杆件不完全理想。实际铰支或固支处存在间隙、摩擦和微小转动，杆件也可能存在初始弯曲，加载轴线与杆件轴线也难以完全重合。这些因素会使实际临界载荷与理论值产生偏差。",
    ),
    (
        "2）观察实验曲线是否理想？",
        "实验曲线不完全理想。曲线中可能出现抖动、平台不明显或局部突变，主要由加载偏心、传感器读数波动、夹具滑移和人工取临界点造成。报告中应取曲线由快速上升转入近似平台阶段的位置作为 Pcr实，而不是取最终 quit point。",
    ),
    (
        "3）哪个因素对实验结果影响最大？",
        "我认为压杆相当长度和约束条件影响最大。因为欧拉公式中 Pcr 与有效长度的平方成反比，长度或约束系数稍有误差，计算出的临界载荷就会明显变化。",
    ),
]
for q, a in answers:
    p = doc.add_paragraph()
    r = p.add_run(q)
    set_run_font(r, size=10.5, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(a)
    set_run_font(r, size=10.5)

add_heading(doc, "六、实验认识与体会", 1)
p = doc.add_paragraph()
r = p.add_run(
    "通过本实验认识到，细长压杆的破坏不一定是强度不足，而可能是稳定性失效。"
    "压杆临界载荷与材料弹性模量、截面惯性矩、杆长和端部约束密切相关，其中约束条件和有效长度影响很大。"
    "实验中还体会到实际装置很难达到理想边界条件，测量长度、安装对中和曲线取点都会影响结果。"
)
set_run_font(r, size=10.5)

doc.save(OUT)
print(OUT)
