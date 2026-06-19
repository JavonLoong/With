# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


BASE_DIR = Path(__file__).resolve().parent
SOURCE = BASE_DIR / "习概期中作业_恢复_2128版.docx"
TMP = BASE_DIR / "习概期中作业_恢复_2128版_圈码脚注内容补强.tmp.docx"
OUTPUT = BASE_DIR / "习概期中作业_恢复_2128版_圈码脚注内容补强版.docx"

CIRCLED_NUMBERS = {
    1: "①",
    2: "②",
    3: "③",
    4: "④",
    5: "⑤",
    6: "⑥",
    7: "⑦",
    8: "⑧",
    9: "⑨",
    10: "⑩",
    11: "⑪",
    12: "⑫",
}

SUPPLEMENTS = [
    (
        "摘 要：",
        " 为避免只停留在概念层面，本文也将结合《弱鸟如何先飞》《滴水穿石的启示》《为官一任，造福一方》《对闽东经济发展的思考》等具体篇章，说明书中如何把人民立场落实为干部作风、产业选择和长期治理方法。",
    ),
    (
        "正因如此，我选择将《摆脱贫困》这本书作为切入点。",
        " 具体看，《弱鸟如何先飞》讨论的是贫困地区如何摆脱精神和思路上的被动，《滴水穿石的启示》强调的是长期积累而不是短期突击；这两篇使“摆脱贫困”不再只是收入指标，而是观念、组织和发展路径同时改变的过程。",
    ),
    (
        "在这本书中，一条清晰的价值主线贯穿始终",
        " 特别是《为官一任，造福一方》所处理的，不是干部个人道德评价，而是权力如何通过修路、办水电、抓教育、整顿作风等具体事务转化为群众能够感受到的公共利益。",
    ),
    (
        "我国扶贫政策从“粗放式灌溉”走向“精准式滴灌”",
        " 这种思路与《对闽东经济发展的思考》中从闽东山区、沿海、资源和产业基础出发判断发展路子的写法相通：先把地方真实条件摸清楚，再谈资源投向和政策安排。",
    ),
    (
        "这就必然引申出另一个关于治理工具的重要命题",
        " 放到《摆脱贫困》的文本中看，“四下基层”不是一个孤立口号，而是把信访接待、现场办公、调查研究、政策宣传同群众的具体困难连接起来；干部只有进入现场，才可能知道哪些问题是材料里看不出来的。",
    ),
    (
        "从历史宏观维度来看，脱贫攻坚不仅是中国摆脱贫穷的面貌改变",
        " 《走一条发展大农业的路子》《建设好贫困地区的精神文明》等篇也说明，脱贫之后能否稳定发展，取决于产业基础、精神文明和公共服务是否持续积累，而不是某一年收入数字的短期达标。",
    ),
    (
        "《摆脱贫困》中虽载有诸多特定时空下的政策细节",
        " 因此，文章所说的基层治理逻辑，并不是从外部套一个现代化理论，而是从《摆脱贫困》内部的篇章结构中读出来的：先认识贫困，再改变干部作风，再选择适合地方的产业和组织方式。",
    ),
]


def set_run_font(run, font_name="宋体", size=12):
    run.font.name = font_name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = etree.Element(qn("w:rFonts"))
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def add_supplements():
    doc = Document(SOURCE)
    applied = []
    for needle, addition in SUPPLEMENTS:
        matched = False
        for paragraph in doc.paragraphs:
            if needle in paragraph.text:
                if addition.strip() not in paragraph.text:
                    run = paragraph.add_run(addition)
                    set_run_font(run)
                applied.append(needle)
                matched = True
                break
        if not matched:
            raise RuntimeError(f"没有找到要补充的段落：{needle}")
    doc.save(TMP)
    return applied


def get_or_create_rpr(run):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rpr = run.find(f"{{{w_ns}}}rPr")
    if rpr is None:
        rpr = etree.Element(f"{{{w_ns}}}rPr")
        run.insert(0, rpr)
    return rpr


def make_visible_mark_run(ref_run, note_id, w_ns):
    mark = CIRCLED_NUMBERS.get(note_id, str(note_id))
    mark_run = etree.Element(f"{{{w_ns}}}r")
    rpr = etree.SubElement(mark_run, f"{{{w_ns}}}rPr")
    fonts = etree.SubElement(rpr, f"{{{w_ns}}}rFonts")
    fonts.set(f"{{{w_ns}}}eastAsia", "宋体")
    fonts.set(f"{{{w_ns}}}ascii", "Times New Roman")
    fonts.set(f"{{{w_ns}}}hAnsi", "Times New Roman")
    vert = etree.SubElement(rpr, f"{{{w_ns}}}vertAlign")
    vert.set(f"{{{w_ns}}}val", "superscript")
    sz = etree.SubElement(rpr, f"{{{w_ns}}}sz")
    sz.set(f"{{{w_ns}}}val", "16")
    sz_cs = etree.SubElement(rpr, f"{{{w_ns}}}szCs")
    sz_cs.set(f"{{{w_ns}}}val", "16")
    text = etree.SubElement(mark_run, f"{{{w_ns}}}t")
    text.text = mark
    return mark_run


def patch_document_circled_refs(root):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    refs = root.xpath("//w:footnoteReference", namespaces=ns)
    for ref in refs:
        note_id = int(ref.get(f"{{{w_ns}}}id"))
        ref.set(f"{{{w_ns}}}customMarkFollows", "1")
        ref_run = ref.getparent()
        rpr = get_or_create_rpr(ref_run)
        if not rpr.xpath("./w:vertAlign[@w:val='superscript']", namespaces=ns):
            vert = etree.SubElement(rpr, f"{{{w_ns}}}vertAlign")
            vert.set(f"{{{w_ns}}}val", "superscript")
        parent = ref_run.getparent()
        index = parent.index(ref_run)

        next_el = parent[index + 1] if index + 1 < len(parent) else None
        expected = CIRCLED_NUMBERS.get(note_id, str(note_id))
        next_text = "".join(next_el.xpath(".//w:t/text()", namespaces=ns)) if next_el is not None else ""
        if next_text != expected:
            parent.insert(index + 1, make_visible_mark_run(ref_run, note_id, w_ns))


def patch_footnotes_circled(root):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    for fn in root.xpath("//w:footnote[not(@w:type)]", namespaces=ns):
        note_id = int(fn.get(f"{{{w_ns}}}id"))
        mark = CIRCLED_NUMBERS.get(note_id, str(note_id))
        ref = fn.xpath(".//w:footnoteRef", namespaces=ns)
        if ref:
            ref_run = ref[0].getparent()
            for child in list(ref_run):
                if child.tag == f"{{{w_ns}}}footnoteRef":
                    ref_run.remove(child)
            if not ref_run.xpath("./w:t", namespaces=ns):
                t = etree.SubElement(ref_run, f"{{{w_ns}}}t")
                t.text = mark
            else:
                ref_run.xpath("./w:t", namespaces=ns)[0].text = mark
            rpr = get_or_create_rpr(ref_run)
            if not rpr.xpath("./w:vertAlign[@w:val='superscript']", namespaces=ns):
                vert = etree.SubElement(rpr, f"{{{w_ns}}}vertAlign")
                vert.set(f"{{{w_ns}}}val", "superscript")
        else:
            first_p = fn.xpath("./w:p", namespaces=ns)[0]
            first_p.insert(0, make_visible_mark_run(first_p, note_id, w_ns))

        # Remove leading spaces created by the previous generator.
        for t in fn.xpath(".//w:t", namespaces=ns):
            if t.text and t.text.startswith(" "):
                t.text = t.text.lstrip()
                break


def patch_footnote_texts(root):
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns}
    for fn in root.xpath("//w:footnote[not(@w:type)]", namespaces=ns):
        note_id = int(fn.get(f"{{{w_ns}}}id"))
        if note_id == 2:
            texts = fn.xpath(".//w:t", namespaces=ns)
            full = "".join(t.text or "" for t in texts)
            appendix = "文中涉及《弱鸟如何先飞》《滴水穿石的启示》《为官一任，造福一方》《对闽东经济发展的思考》等篇。"
            if appendix not in full:
                last_text = texts[-1]
                if last_text.text and not last_text.text.endswith("。"):
                    last_text.text += "。"
                last_text.text = (last_text.text or "") + appendix


def patch_docx():
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    with tempfile.TemporaryDirectory() as td:
        tmp_copy = Path(td) / "work.docx"
        shutil.copyfile(TMP, tmp_copy)
        with zipfile.ZipFile(tmp_copy, "r") as zin, zipfile.ZipFile(OUTPUT, "w", zipfile.ZIP_DEFLATED) as zout:
            doc_xml = etree.fromstring(zin.read("word/document.xml"))
            patch_document_circled_refs(doc_xml)

            footnotes_xml = etree.fromstring(zin.read("word/footnotes.xml"))
            patch_footnotes_circled(footnotes_xml)
            patch_footnote_texts(footnotes_xml)

            for item in zin.infolist():
                if item.filename in {"word/document.xml", "word/footnotes.xml"}:
                    continue
                zout.writestr(item, zin.read(item.filename))
            zout.writestr(
                "word/document.xml",
                etree.tostring(doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
            zout.writestr(
                "word/footnotes.xml",
                etree.tostring(footnotes_xml, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
    TMP.unlink(missing_ok=True)


def main():
    applied = add_supplements()
    patch_docx()
    print(OUTPUT)
    print("supplements", len(applied))


if __name__ == "__main__":
    main()
