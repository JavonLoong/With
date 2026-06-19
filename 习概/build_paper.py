# -*- coding: utf-8 -*-
from pathlib import Path
import re
import shutil
import tempfile
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "习概期中作业_课程论文初稿.docx"
MD_PATH = OUT_DIR / "习概期中作业_课程论文初稿.md"

TITLE = "从《摆脱贫困》看以人民为中心的发展思想"
SUBTITLE = "——兼论精准扶贫的基层治理逻辑"

FOOTNOTES = {
    1: "本书编写组：《习近平新时代中国特色社会主义思想概论（2023年版）》，高等教育出版社、人民出版社，2023年，第4章“坚持以人民为中心”。",
    2: "习近平：《摆脱贫困》，福建人民出版社，1992年。",
    3: "陈振明：《〈摆脱贫困〉中的地方治理思想研究》。",
    4: "汪三贵、郭子豪：《论中国的精准扶贫》，《贵州社会科学》2015年第5期。",
    5: "葛志军、邢成举：《精准扶贫：内涵、实践困境及其原因阐释——基于宁夏银川两个村庄的调查》，《贵州社会科学》2015年第5期。",
    6: "邓维杰：《精准扶贫的难点、对策与路径选择》。",
    7: "申坤：《新时代背景下习近平调查研究重要论述的价值意义》，《山东干部函授大学学报》2020年第3期。",
    8: "杨明伟：《百年奋斗史中的摆脱贫困迈向共同富裕》，2021年。",
    9: "胡蝶：《扶贫必扶智：教育精准扶贫是摆脱贫困的内生动力》。",
    10: "习近平：《在决战决胜脱贫攻坚座谈会上的讲话》，2020年3月6日。",
}

SECTIONS = [
    (
        None,
        [
            [
                "摘  要：",
                "《摆脱贫困》写于习近平同志在福建宁德工作之后，讨论的是闽东如何摆脱贫困的问题，但其中呈现的并不只是某一地区的发展经验。它把人民立场、群众路线、调查研究、干部作风和发展路径放在同一个治理场景中加以思考。本文结合“习近平新时代中国特色社会主义思想概论”课程内容，尝试从以人民为中心的发展思想出发，分析《摆脱贫困》中的基层治理逻辑，并联系精准扶贫的实践说明：脱贫攻坚之所以能够成为中国式现代化进程中的重要经验，关键不在于单一政策工具，而在于把人民的实际生活作为政策出发点和检验标准。"
            ],
            [
                "关键词：",
                "《摆脱贫困》；以人民为中心；精准扶贫；基层治理；共同富裕"
            ],
        ],
    ),
    (
        "一、问题的提出：为什么从“摆脱贫困”理解以人民为中心",
        [
            [
                "在学习“坚持以人民为中心”这一章时，我最直接的感受是，这个命题如果只停留在概念层面，容易被理解成一句原则性的表达；但一旦放到贫困治理中，它就变得很具体：谁被看见，谁被帮助，谁的发展机会被真正改善，都是可以检验的。教材把人民立场视为习近平新时代中国特色社会主义思想的根本政治立场，强调发展为了人民、发展依靠人民、发展成果由人民共享。",
                1,
                "这实际上提示我们，评价一种发展政策，不能只看宏观指标，也要看普通人的获得感是否真实。"
            ],
            [
                "因此，我选择从《摆脱贫困》切入，不只是因为它与精准扶贫、共同富裕等课程内容联系紧密，更因为它提供了一个观察基层治理的入口。贫困并不是抽象数字，它常常体现为交通、教育、观念、产业、干部作风等多方面问题的叠加。要改变这种状态，仅仅靠口号或者一次性的救济是不够的，必须把发展问题、组织问题和人的主体性问题放在一起处理。"
            ],
            [
                "在这个意义上，摆脱贫困可以被看作以人民为中心发展思想的一个实践样本。它不是把人民当作被动接受政策的对象，而是要求干部深入群众、理解群众、发动群众，并在具体工作中把群众利益作为判断得失的标准。"
            ],
        ],
    ),
    (
        "二、《摆脱贫困》中的价值线索：为民造福不是空泛口号",
        [
            [
                "《摆脱贫困》中有一个很鲜明的价值取向，就是把干部的责任同一方百姓的生活联系起来。书中关于“为官一场，造福一方”的论述，表达的并不是简单的道德要求，而是一种治理观：干部手中的权力如果不能转化为群众生活的改善，就失去了应有的公共意义。",
                2,
                "这种理解与课程中讲到的人民立场是相通的。人民不是政策文件中的抽象名词，而是一个个具体的人，是需要道路、教育、就业、医疗和发展机会的人。"
            ],
            [
                "从地方治理角度看，《摆脱贫困》强调的“为民造福”至少有两层含义。第一，干部必须有责任意识，不能把贫困地区的发展困难简单归因于群众落后、条件不好，更不能用形式主义的办法应付任务。第二，干部又不能代替群众包办一切，而要通过组织、引导和服务，把群众自身的积极性调动起来。陈振明对《摆脱贫困》地方治理思想的研究也指出，其中贯穿着人民群众根本利益至上的公共价值追求。",
                3,
                "这说明，人民立场并不是停留在态度上，而要落实为治理行为。"
            ],
            [
                "我觉得这一点尤其值得注意。很多时候，我们谈扶贫容易只谈“给了什么”，比如资金、项目、物资；但《摆脱贫困》更重视“怎样让一个地方真正发展起来”。如果一个地方只是短期增加收入，却没有形成稳定产业、公共服务和人的能力提升，那么贫困问题可能还会以新的形式出现。也正因为如此，摆脱贫困与共同富裕之间不是两个割裂目标，而是一个连续过程：前者解决底线问题，后者继续回答发展是否更均衡、更充分。"
            ],
        ],
    ),
    (
        "三、精准扶贫的治理逻辑：把“人民”具体到每一户、每一项需求",
        [
            [
                "精准扶贫之所以重要，正在于它试图克服传统扶贫中“粗放”和“脱靶”的问题。汪三贵、郭子豪在研究中指出，精准扶贫面临精准识别、精准扶持和精准考核三个方面的困难，同时需要通过机制创新来保障实际效果。",
                4,
                "这与课程所强调的实事求是、问题导向是一致的。贫困人口在哪里、因为什么致贫、需要什么样的帮扶，如果这些问题没有弄清楚，扶贫政策就可能变成平均用力，甚至出现真正困难的人没有被扶到的情况。"
            ],
            [
                "精准扶贫的关键，不只是“精准”两个字好听，而是它把人民立场具体化了。所谓以人民为中心，在扶贫工作中就不能只是面向一个地区、一个平均数，而要落到具体村庄、具体家庭和具体原因。比如有的家庭是因病致贫，有的是缺少产业，有的是教育机会不足，有的是交通条件限制。不同原因需要不同办法，这就是精准帮扶的现实意义。葛志军、邢成举对宁夏银川两个村庄的调查也提到，精准扶贫包含精准识别、帮扶、管理和考核等环节，但地方实践中会遇到贫困户参与不足、政策缺乏灵活性、资金有限等困境。",
                5,
                "这些困境提醒我们，好的政策理念必须通过细致的基层执行才能真正落地。"
            ],
            [
                "同时，精准扶贫并不是没有问题的理想过程。邓维杰在讨论精准扶贫难点时提到，实践中可能出现规模排斥、区域排斥和识别排斥等问题。",
                6,
                "我认为这恰恰说明，以人民为中心不是一句完成时的话，而是一种持续校正政策偏差的方法。如果识别贫困户的过程缺少公开、公平和群众参与，那么政策名称再精准，也可能偏离真正需要帮助的人。反过来说，越是基层工作复杂，越需要把群众路线、民主评议、信息公开和监督机制结合起来。"
            ],
            [
                "从这个角度看，精准扶贫的治理逻辑可以概括为三点：第一，用调查研究代替想当然；第二，用分类施策代替一刀切；第三，用群众实际受益代替形式上的完成任务。它的重点不是制造一套复杂表格，而是让政策能够贴近真实生活。"
            ],
        ],
    ),
    (
        "四、调查研究与群众路线：基层治理不能只在材料里完成",
        [
            [
                "《摆脱贫困》给我的另一个启发，是贫困治理必须从调查研究开始。申坤在文章中提到，习近平同志在宁德工作时曾倡导“四下基层”，并强调调查研究要从群众中来、到群众中去。",
                7,
                "这与我们在课程中学习的群众路线是一致的。群众路线不是简单地“听取意见”，而是要求干部真正进入现场，理解问题发生的条件，再把群众经验转化为政策判断。"
            ],
            [
                "如果离开调查研究，基层治理很容易出现两种偏差：一种是上级想象基层，另一种是基层迎合指标。前者会使政策设计脱离现实，后者会使执行过程重材料、轻效果。贫困治理尤其不能这样。一个家庭是否需要帮扶，不能只看表格上某个收入数字；一个村庄是否有发展基础，也不能只看短期项目是否建成。只有走到农户家里、田间地头和村庄公共空间，才能看到那些材料里不容易呈现的问题。"
            ],
            [
                "这也是我认为《摆脱贫困》仍然有现实意义的地方。它没有把贫困地区写成等待拯救的对象，而是反复讨论干部如何认识地方、如何联系群众、如何找到适合本地的发展路子。这样的思路对于今天的乡村振兴同样重要。脱贫之后，基层治理的任务没有消失，而是转向防止返贫、提升公共服务、培育产业和改善治理能力等更长期的问题。"
            ],
        ],
    ),
    (
        "五、从脱贫攻坚到共同富裕：发展成果怎样更稳定地留在人民生活中",
        [
            [
                "脱贫攻坚的全面胜利，标志着我国在解决绝对贫困问题上取得了历史性成就。杨明伟在回顾百年奋斗史时指出，摆脱贫困、迈向共同富裕是中国人民近代以来的重要追求，也是实现民族复兴的重要内容。",
                8,
                "但从课程视角看，脱贫攻坚并不是终点。中国式现代化强调全体人民共同富裕，这意味着发展不能只追求总量增长，还要持续关注区域差距、城乡差距和不同群体的发展能力。"
            ],
            [
                "共同富裕不是平均主义，也不是短期分配口号，而是一个需要高质量发展支撑的长期过程。脱贫攻坚解决了绝对贫困问题，但一些地区在产业基础、教育水平、公共服务和就业机会方面仍然相对薄弱。胡蝶关于教育精准扶贫的文章提出，教育是阻断代际贫困、提升贫困地区内生动力的重要手段。",
                9,
                "这一点对理解共同富裕很有帮助。真正稳定的脱贫，不只是当下收入达标，还包括下一代拥有更好的发展机会。"
            ],
            [
                "因此，从《摆脱贫困》到精准扶贫，再到乡村振兴和共同富裕，可以看到一条比较清楚的逻辑：第一步是让困难群众摆脱基本生活困境，第二步是通过产业、教育和公共服务增强发展能力，第三步是在更高水平上实现发展成果共享。习近平总书记在决战决胜脱贫攻坚座谈会上的讲话中强调，要克服新冠肺炎疫情影响，凝心聚力打赢脱贫攻坚战。",
                10,
                "这种部署体现了一个重要判断：越是面对复杂困难，越要把人民生活放在政策优先位置。"
            ],
            [
                "对大学生来说，这个问题也不只是宏观叙事。我们讨论基层治理、乡村振兴和共同富裕时，不能只把它们当成考试概念，而要看到背后的真实社会关系。比如，一个地方能不能留住年轻人，乡村学校有没有好老师，基层干部有没有足够能力和动力，公共资源能不能公平配置，这些都关系到人民中心的发展思想能否落地。"
            ],
        ],
    ),
    (
        "结语",
        [
            [
                "总体来看，《摆脱贫困》中的许多论述虽然产生于特定历史和地区背景，却包含着理解新时代中国治理的重要线索。它告诉我们，以人民为中心不是抽象价值宣示，而要体现在干部作风、政策设计、调查研究、群众参与和发展成效之中。精准扶贫之所以能够成为脱贫攻坚的重要方法，也正在于它把“人民”从总体概念还原为一个个具体家庭和具体需求。"
            ],
            [
                "当然，贫困治理的复杂性也提醒我们，任何政策都需要在实践中不断校正。真正的人民立场，不是把政策说得完美，而是愿意承认问题、进入现场、听取群众意见，并根据现实情况调整办法。由此看，摆脱贫困与共同富裕之间的联系，不只是目标上的递进，更是治理方法上的延续：坚持党的领导，坚持人民主体地位，坚持实事求是，最终让发展成果更稳定、更公平地转化为人民的现实生活。"
            ],
        ],
    ),
]


def set_run_font(run, font_name="宋体", size=None, bold=None):
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def apply_para_format(paragraph, first_line=True, after=8, line=1.333, justify=True):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(24)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_footnote_reference(paragraph, note_id):
    run_element = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "FootnoteReference")
    rpr.append(style)
    run_element.append(rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(note_id))
    run_element.append(ref)
    paragraph._p.append(run_element)


def add_mixed_paragraph(doc, parts):
    p = doc.add_paragraph()
    apply_para_format(p)
    for part in parts:
        if isinstance(part, int):
            add_footnote_reference(p, part)
        else:
            run = p.add_run(part)
            set_run_font(run, "宋体", 12)
    return p


def build_docx_base(path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run(TITLE)
    set_run_font(r, "黑体", 16, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run(SUBTITLE)
    set_run_font(r, "楷体", 12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    r = meta.add_run("姓名：________    学号：________    班级：________")
    set_run_font(r, "宋体", 11)

    for heading, paragraphs in SECTIONS:
        if heading:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(heading)
            set_run_font(r, "黑体", 13, True)
        for parts in paragraphs:
            add_mixed_paragraph(doc, parts)

    doc.save(path)


def make_footnotes_xml():
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xml_ns = "http://www.w3.org/XML/1998/namespace"
    footnotes = etree.Element(f"{{{w}}}footnotes", nsmap={"w": w})

    sep = etree.SubElement(footnotes, f"{{{w}}}footnote")
    sep.set(f"{{{w}}}type", "separator")
    sep.set(f"{{{w}}}id", "-1")
    p = etree.SubElement(sep, f"{{{w}}}p")
    r = etree.SubElement(p, f"{{{w}}}r")
    etree.SubElement(r, f"{{{w}}}separator")

    cont = etree.SubElement(footnotes, f"{{{w}}}footnote")
    cont.set(f"{{{w}}}type", "continuationSeparator")
    cont.set(f"{{{w}}}id", "0")
    p = etree.SubElement(cont, f"{{{w}}}p")
    r = etree.SubElement(p, f"{{{w}}}r")
    etree.SubElement(r, f"{{{w}}}continuationSeparator")

    for note_id, text in FOOTNOTES.items():
        fn = etree.SubElement(footnotes, f"{{{w}}}footnote")
        fn.set(f"{{{w}}}id", str(note_id))
        p = etree.SubElement(fn, f"{{{w}}}p")
        ppr = etree.SubElement(p, f"{{{w}}}pPr")
        pstyle = etree.SubElement(ppr, f"{{{w}}}pStyle")
        pstyle.set(f"{{{w}}}val", "FootnoteText")
        r = etree.SubElement(p, f"{{{w}}}r")
        rpr = etree.SubElement(r, f"{{{w}}}rPr")
        rstyle = etree.SubElement(rpr, f"{{{w}}}rStyle")
        rstyle.set(f"{{{w}}}val", "FootnoteReference")
        etree.SubElement(r, f"{{{w}}}footnoteRef")
        text_run = etree.SubElement(p, f"{{{w}}}r")
        text_rpr = etree.SubElement(text_run, f"{{{w}}}rPr")
        fonts = etree.SubElement(text_rpr, f"{{{w}}}rFonts")
        fonts.set(f"{{{w}}}eastAsia", "宋体")
        fonts.set(f"{{{w}}}ascii", "Times New Roman")
        fonts.set(f"{{{w}}}hAnsi", "Times New Roman")
        sz = etree.SubElement(text_rpr, f"{{{w}}}sz")
        sz.set(f"{{{w}}}val", "18")
        sz_cs = etree.SubElement(text_rpr, f"{{{w}}}szCs")
        sz_cs.set(f"{{{w}}}val", "18")
        t = etree.SubElement(text_run, f"{{{w}}}t")
        t.set(f"{{{xml_ns}}}space", "preserve")
        t.text = " " + text

    return etree.tostring(footnotes, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_footnotes(src, dst):
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    foot_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp) / "base.docx"
        shutil.copyfile(src, tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
            content_types = etree.fromstring(zin.read("[Content_Types].xml"))
            existing = content_types.xpath(
                "ct:Override[@PartName='/word/footnotes.xml']",
                namespaces={"ct": ct_ns},
            )
            if not existing:
                override = etree.SubElement(content_types, f"{{{ct_ns}}}Override")
                override.set("PartName", "/word/footnotes.xml")
                override.set(
                    "ContentType",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
                )

            rels = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
            if not rels.xpath("rel:Relationship[@Type=$rtype]", namespaces={"rel": rel_ns}, rtype=foot_rel_type):
                ids = []
                for rel in rels.xpath("rel:Relationship", namespaces={"rel": rel_ns}):
                    rid = rel.get("Id", "")
                    match = re.match(r"rId(\d+)$", rid)
                    if match:
                        ids.append(int(match.group(1)))
                rel = etree.SubElement(rels, f"{{{rel_ns}}}Relationship")
                rel.set("Id", f"rId{max(ids or [0]) + 1}")
                rel.set("Type", foot_rel_type)
                rel.set("Target", "footnotes.xml")

            for item in zin.infolist():
                if item.filename in {
                    "[Content_Types].xml",
                    "word/_rels/document.xml.rels",
                    "word/footnotes.xml",
                }:
                    continue
                zout.writestr(item, zin.read(item.filename))

            zout.writestr(
                "[Content_Types].xml",
                etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
            zout.writestr(
                "word/_rels/document.xml.rels",
                etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
            zout.writestr("word/footnotes.xml", make_footnotes_xml())


def parts_to_md(parts):
    out = []
    for part in parts:
        if isinstance(part, int):
            out.append(f"[^{part}]")
        else:
            out.append(part)
    return "".join(out)


def build_markdown(path):
    lines = [
        f"# {TITLE}",
        "",
        SUBTITLE,
        "",
        "姓名：________    学号：________    班级：________",
        "",
    ]
    for heading, paragraphs in SECTIONS:
        if heading:
            lines.extend([f"## {heading}", ""])
        for parts in paragraphs:
            lines.extend([parts_to_md(parts), ""])
    lines.append("## 脚注")
    lines.append("")
    for note_id, text in FOOTNOTES.items():
        lines.append(f"[^{note_id}]: {text}")
    path.write_text("\n".join(lines), encoding="utf-8")


def main():
    with tempfile.TemporaryDirectory() as tmp:
        base = Path(tmp) / "base.docx"
        build_docx_base(base)
        patch_docx_footnotes(base, DOCX_PATH)
    build_markdown(MD_PATH)
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
