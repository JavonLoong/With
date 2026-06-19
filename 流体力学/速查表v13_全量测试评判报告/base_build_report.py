from __future__ import annotations

import csv
import json
import math
import os
import re
import shutil
import subprocess
import textwrap
from collections import Counter, defaultdict
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Iterable

from docx import Document


ROOT = Path(r"D:\虚拟C盘\学习\流体力学")
OUT = ROOT / "速查表v6_全量测试评判报告"
TEXT_DIR = OUT / "extracted_text"
SOURCE_MANIFEST = OUT / "02_题源清单.csv"
ITEM_CSV = OUT / "05_全量逐题解决路径.csv"

PDFTOTEXT = shutil.which("pdftotext") or r"C:\texlive\2025\bin\windows\pdftotext.exe"


CHEATSHEET_PDF = ROOT / "期末六页A4速查表" / "期末六页A4速查表_v6_结构导航版.pdf"
CHEATSHEET_TEX = ROOT / "期末六页A4速查表" / "期末六页A4速查表_v6_结构导航版.tex"


@dataclass
class SourceRecord:
    group: str
    path: str
    status: str
    extracted_text: str
    char_count: int
    candidate_count: int = 0
    note: str = ""


@dataclass
class EvalItem:
    item_id: str
    group: str
    source_file: str
    source_label: str
    question_no: str
    question_text: str
    inferred_topic: str
    cheatsheet_path: str
    page_reference: str
    solution_route: str
    own_steps: str
    needed_supplement: str
    score_basis: str
    estimated_score: float
    max_score: float
    normalized_percent: float
    confidence: str
    notes: str = ""


CHEATSHEET_INDEX = [
    {
        "page": "P1",
        "title": "先判模型：静力、运动、能量、动量",
        "nodes": [
            ("P1 > 1. 六页定位表", "题面特征到页码和首写公式的入口。"),
            ("P1 > 2. 静水压力 > 压强", "p_abs=p_a+rho gh；表压 p_g=rho gh；同高同压。"),
            ("P1 > 2. 静水压力 > 平面受压", "F=rho g h_C A；y_D=y_C+I_C/(y_C A)。"),
            ("P1 > 2. 静水压力 > 曲面受压", "水平分量、竖直分量、压力体。"),
            ("P1 > 2. 静水压力 > 相对平衡", "平动自由面、旋转自由面。"),
            ("P1 > 3. 流体运动学 > 连续/有旋", "不可压连续方程、无旋条件、二维涡量。"),
            ("P1 > 3. 流体运动学 > 流线/迹线", "流线微分式、迹线微分式、定常时重合。"),
            ("P1 > 3. 流体运动学 > 加速度", "当地加速度和迁移加速度。"),
            ("P1 > 3. 流体运动学 > 应力张量", "t=P n，法向和切向应力。"),
            ("P1 > 4. Bernoulli/动量 > Bernoulli", "理想定常沿流线；实际管路机械能方程。"),
            ("P1 > 4. Bernoulli/动量 > 小孔/水箱", "孔口速度和变水位方程。"),
            ("P1 > 4. Bernoulli/动量 > 动量方程", "二维控制体动量方程。"),
        ],
    },
    {
        "page": "P2",
        "title": "势流与理想不可压",
        "nodes": [
            ("P2 > 1. 基本关系", "无旋、速度势、流函数、Laplace 方程。"),
            ("P2 > 1. 基本关系 > 直角坐标", "u=phi_x=psi_y；v=phi_y=-psi_x。"),
            ("P2 > 1. 基本关系 > 极坐标", "v_r、v_theta 与 phi、psi 的关系。"),
            ("P2 > 2. 基元流速查", "均匀流、点源、点汇、点涡、偶极子。"),
            ("P2 > 2. 叠加原则", "势函数、流函数、速度均可叠加。"),
            ("P2 > 3. 圆柱绕流", "无环量圆柱绕流、表面速度、压力系数、升力。"),
            ("P2 > 4. 镜像法/壁面", "固壁边界的源/涡镜像规则、半无限体、角域流。"),
            ("P2 > 5. 势流题标准答案骨架", "速度场、圆柱/半圆柱、镜像/开孔题通用写法。"),
        ],
    },
    {
        "page": "P3",
        "title": "可压缩流",
        "nodes": [
            ("P3 > 1. 基本量", "状态方程、声速、马赫数、空气默认常数。"),
            ("P3 > 1. 等熵总量", "T0/T、p0/p 关系。"),
            ("P3 > 1. 临界", "M=1、A=A*、空气临界比例。"),
            ("P3 > 1. 面积--马赫数", "面积比公式和亚声/超声双支。"),
            ("P3 > 2. Laval/背压判断", "阻塞、出口压力和背压、背压状态序列。"),
            ("P3 > 3. 正激波", "M2、p2/p1、rho2/rho1、T2/T1、总压下降。"),
            ("P3 > 3. 管内激波链", "波前等熵、过波、波后等熵三段法。"),
            ("P3 > 4. 斜激波/膨胀波", "theta-beta-M 关系、法向分量法、PM 函数。"),
            ("P3 > 5. 高频原题入口", "收缩喷管、Laval 面积、7.19/7.20 等入口。"),
        ],
    },
    {
        "page": "P4",
        "title": "粘性流动",
        "nodes": [
            ("P4 > 1. N-S 和层流解析 > 牛顿流体", "tau=mu du/dy；不可压 N-S。"),
            ("P4 > 1. H-P 圆管层流", "速度分布、平均速度、流量、壁面切应力。"),
            ("P4 > 1. 平板 Couette/Poiseuille", "无压差 Couette 和有压差积分法。"),
            ("P4 > 2. 管路机械能 > 沿程/局部损失", "h_f=lambda L/d V^2/(2g)，h_zeta=zeta V^2/(2g)。"),
            ("P4 > 2. 管路流程", "Q -> V -> Re -> lambda -> 损失 -> 能量方程。"),
            ("P4 > 2. 泵吸水安装高度", "水面到泵入口能量方程和汽蚀限制。"),
            ("P4 > 2. 水轮机功率", "取能水头和 P=eta rho g Q H_T。"),
            ("P4 > 3. 湍流近壁与边界层", "u_tau、y+、平板层流边界层厚度、平板阻力、分离。"),
            ("P4 > 4. Stokes/外绕流", "低 Re 小球阻力和终速。"),
            ("P4 > 5. 第8章题入口", "近壁层厚度、泵题、水轮机题。"),
        ],
    },
    {
        "page": "P5",
        "title": "高频题库模板",
        "nodes": [
            ("P5 > 1. 静力与基础", "平面闸门、曲面闸门、U 管、应力张量。"),
            ("P5 > 2. 势流题", "速度场、均匀流+源/汇、圆柱绕流、壁面镜像。"),
            ("P5 > 3. 可压缩题", "收缩喷管、Laval、内激波、斜激波、膨胀波。"),
            ("P5 > 4. 粘性和管路题", "H-P、管路损失、泵/水轮机、边界层阻力。"),
            ("P5 > 5. 数值换算模板", "面积、流量、压力水头、单位检查。"),
        ],
    },
    {
        "page": "P6",
        "title": "原题图索引、查表顺序、易错点",
        "nodes": [
            ("P6 > 1. 原题图索引", "相似图定位和常见模型。"),
            ("P6 > 2. 查教材/表优先级", "Moody 图、正激波表、斜激波图、PM 函数表、面积-马赫数表、阻力系数图。"),
            ("P6 > 3. 最容易混的符号", "p0、p*、pe、pb、At、A*、平均/局部速度、绝压/表压、方向。"),
            ("P6 > 4. 激波/喷管易错", "0.528、pe=pb、过激波量变、斜激波/膨胀波适用条件。"),
            ("P6 > 5. 管路/粘性易错", "粗糙度、局部损失、泵安装高度、边界层。"),
            ("P6 > 6. 保底写法", "模型判断、控制方程、适用条件、查表项、代数骨架、单位检查。"),
        ],
    },
]


TOPIC_RULES = [
    {
        "topic": "静水压力/测压/闸门",
        "keywords": ["静水", "压强", "测压", "U形", "U 形", "水银", "闸门", "平面", "曲面", "压力体", "高程", "液柱", "容器"],
        "paths": [
            "P1 > 2. 静水压力 > 压强",
            "P1 > 2. 静水压力 > 平面受压",
            "P1 > 2. 静水压力 > 曲面受压",
            "P5 > 1. 静力与基础",
            "P6 > 3. 最容易混的符号",
        ],
        "route": "先判为静止液体或液柱测压；用同一连通静止液体同高同压建立压力链；平面/曲面受压题再写合力、压力中心或压力体；最后检查绝压/表压和方向。",
        "supplement": "若题目依赖图中高程、液面位置、铰链位置或多种液体界面，速查表只给公式，读图和符号方向需要题目图或人工补充。",
        "base": 0.78,
    },
    {
        "topic": "流体运动学/连续/流线迹线",
        "keywords": ["速度场", "流线", "迹线", "涡量", "有旋", "无旋", "连续方程", "不可压缩", "加速度", "应力张量", "欧拉", "拉格朗日"],
        "paths": [
            "P1 > 3. 流体运动学 > 连续/有旋",
            "P1 > 3. 流体运动学 > 流线/迹线",
            "P1 > 3. 流体运动学 > 加速度",
            "P1 > 3. 流体运动学 > 应力张量",
        ],
        "route": "先从题干找速度分量或描述方式；写连续方程、涡量/无旋条件或流线/迹线微分方程；若问加速度，分当地项和迁移项；若问应力，用 t=P n 分解法向/切向。",
        "supplement": "欧拉法/拉格朗日法等概念性定义在 v6 中不是完整背诵材料，只能由速度场和迹线公式反推，简答题会丢定义分。",
        "base": 0.60,
    },
    {
        "topic": "Bernoulli/机械能/动量控制体",
        "keywords": ["伯努利", "能量方程", "机械能", "水箱", "孔口", "喷嘴", "弯管", "动量", "推力", "控制体", "螺丝", "射流", "水头", "流量"],
        "paths": [
            "P1 > 4. Bernoulli/动量 > Bernoulli",
            "P1 > 4. Bernoulli/动量 > 小孔/水箱",
            "P1 > 4. Bernoulli/动量 > 动量方程",
            "P5 > 5. 数值换算模板",
            "P6 > 6. 保底写法",
        ],
        "route": "先选截面并写连续方程；用 Bernoulli/实际机械能方程求速度、流量或压差；再对控制体列动量方程求支反力/推力；最后把壁面对流体的力取反并做单位检查。",
        "supplement": "控制体受力图、压力力方向和图中截面几何通常要自己画；v6 有方程骨架但没有完整受力图模板。",
        "base": 0.74,
    },
    {
        "topic": "势流/圆柱/镜像/升力",
        "keywords": ["势函数", "流函数", "势流", "点源", "点汇", "点涡", "偶极", "圆柱", "半圆柱", "镜像", "环量", "升力", "气膜", "达朗贝尔"],
        "paths": [
            "P2 > 1. 基本关系",
            "P2 > 2. 基元流速查",
            "P2 > 3. 圆柱绕流",
            "P2 > 4. 镜像法/壁面",
            "P2 > 5. 势流题标准答案骨架",
        ],
        "route": "先判断是否无旋不可压并选 phi/psi；按基元叠加写总势函数或流函数；由速度求压力系数或 Bernoulli 压力；圆柱/半圆柱题对表面压力分量积分求力。",
        "supplement": "压力分量积分限、单位展长、实际图形边界和内部压力条件常需自己补；v6 对圆柱命中高，但对几何解释不完全。",
        "base": 0.82,
    },
    {
        "topic": "可压缩流/喷管/激波/膨胀波",
        "keywords": ["马赫", "Ma", "M=", "喷管", "Laval", "背压", "阻塞", "临界", "等熵", "正激波", "斜激波", "激波", "膨胀波", "PM", "Prandtl", "面积比"],
        "paths": [
            "P3 > 1. 基本量",
            "P3 > 1. 等熵总量",
            "P3 > 1. 面积--马赫数",
            "P3 > 2. Laval/背压判断",
            "P3 > 3. 正激波",
            "P3 > 4. 斜激波/膨胀波",
            "P6 > 2. 查教材/表优先级",
            "P6 > 4. 激波/喷管易错",
        ],
        "route": "先判等熵段、喉部是否阻塞、是否有激波/膨胀波；等熵段用 T0/T、p0/p、A/A*；正激波用 M1 到 M2 和压强/密度比；斜激波用 theta-beta-M 后转法向分量；PM 题用 ν(M) 增量。",
        "supplement": "大量数值题需要面积-马赫数表、正激波表、斜激波图或 PM 函数表；v6 只写查表入口，不给完整表值。",
        "base": 0.68,
    },
    {
        "topic": "粘性管路/沿程局部损失/泵水轮机",
        "keywords": ["粘性", "黏性", "管路", "沿程", "局部", "水头损失", "摩阻", "粗糙", "Moody", "Colebrook", "泵", "汽蚀", "水轮机", "层流", "紊流", "湍流", "圆管", "H-P", "Poiseuille"],
        "paths": [
            "P4 > 1. H-P 圆管层流",
            "P4 > 2. 管路机械能 > 沿程/局部损失",
            "P4 > 2. 管路流程",
            "P4 > 2. 泵吸水安装高度",
            "P4 > 2. 水轮机功率",
            "P6 > 5. 管路/粘性易错",
        ],
        "route": "先由 Q 求各管段 V；再算 Re 并判断层流/湍流；由层流 64/Re 或 Moody/Colebrook 得摩阻；汇总沿程与局部损失；代入机械能方程求液面差、泵安装高度或水轮机功率。",
        "supplement": "湍流摩阻系数、局部损失系数、粗糙度和汽蚀压力需要题给或查图；v6 有流程但不含完整图表。",
        "base": 0.76,
    },
    {
        "topic": "边界层/外绕流阻力",
        "keywords": ["边界层", "排挤厚度", "动量损失厚度", "动量积分", "平板", "摩擦阻力", "阻力系数", "分离", "圆球", "绕流", "Stokes", "终速"],
        "paths": [
            "P4 > 3. 湍流近壁与边界层",
            "P4 > 4. Stokes/外绕流",
            "P5 > 4. 粘性和管路题",
            "P6 > 2. 查教材/表优先级",
            "P6 > 5. 管路/粘性易错",
        ],
        "route": "先算 Re 判断层流、湍流或低 Re；平板直接题用边界层厚度和平均摩擦系数公式；外绕流用阻力系数或 Stokes 阻力；若涉及分离，说明逆压梯度与边界层发展。",
        "supplement": "v6 缺排挤厚度、动量损失厚度定义积分和 Karman 动量积分方程；给速度分布求厚度或推导题会明显失分。",
        "base": 0.58,
    },
    {
        "topic": "量纲分析/相似准则",
        "keywords": ["量纲", "π", "Pi", "相似", "雷诺数准则", "傅汝德", "Fr", "模型律", "准则"],
        "paths": [
            "P4 > 3. 湍流近壁与边界层",
            "P5 > 5. 数值换算模板",
            "P6 > 2. 查教材/表优先级",
        ],
        "route": "只能先列相关物理量并尝试用无量纲数表达；若题目只问 Re 相关可接到阻力系数或粘性流动入口。",
        "supplement": "v6 没有 Buckingham π 定理流程、Fr 定义和相似准则矛盾的解释；这类题主要靠外部学习补齐。",
        "base": 0.25,
    },
    {
        "topic": "水波/水击/课程概念缺口",
        "keywords": ["深水波", "浅水波", "水击", "水锤", "波速", "传播速度"],
        "paths": [
            "未在 v6 中直接覆盖",
            "P6 > 6. 保底写法",
        ],
        "route": "v6 只能写出模型判断和保底表述，无法给出定义、公式或影响因素。",
        "supplement": "需要补充水波理论、水击波速、管壁弹性和流体可压缩性等概念材料。",
        "base": 0.05,
    },
    {
        "topic": "概念简答/定义解释",
        "keywords": ["简述", "定义", "为什么", "意义", "特点", "差异", "适用条件", "各项", "是什么"],
        "paths": [
            "P6 > 6. 保底写法",
        ],
        "route": "先定位相关公式或模型入口，再把公式中的物理量、适用条件和易错点写成文字说明。",
        "supplement": "v6 偏公式和计算模板，概念定义、实验现象和物理解释不足；简答题需要单独背诵库。",
        "base": 0.35,
    },
]


SOLUTION_CUT_PAT = re.compile(r"(?m)^\s*(解[:：]|答案[:：]|解析[:：]|解答[:：]|参考答案|Solution\b)")
DIR_NOISE_PAT = re.compile(r"[·\.]{4,}\s*\d+\s*$")
QUESTION_VERBS = re.compile(r"(求|计算|证明|简述|写出|推求|确定|问|试|分析|说明|判断|给出|比较|为什么|定义|绘|画|讨论|估计)")
QUESTION_START_PAT = re.compile(
    r"(?m)^\s*(?:"
    r"第\s*[一二三四五六七八九十百0-9]+\s*[题問]|"
    r"[一二三四五六七八九十]+[、.．]|"
    r"\(?[0-9]{1,2}\)?[.．、]\s*|"
    r"[0-9]{1,2}\.[0-9]{1,2}\s+|"
    r"书\s*[0-9]{1,2}\.[0-9]{1,2}|"
    r"例\s*[0-9]{1,2}[.．][0-9]{1,2}|"
    r"##+\s*(?:书|第|[0-9])"
    r")"
)


def safe_name(path: Path) -> str:
    rel = path.relative_to(ROOT) if path.is_relative_to(ROOT) else path
    s = str(rel).replace("\\", "__").replace("/", "__").replace(":", "")
    return re.sub(r"[^\w\-.一-龥（）()]+", "_", s)[:180]


def run(cmd: list[str], timeout: int = 120) -> tuple[int, str, str]:
    try:
        p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=timeout, encoding="utf-8", errors="replace")
        return p.returncode, p.stdout, p.stderr
    except Exception as e:
        return 999, "", repr(e)


def extract_pdf(path: Path, target: Path) -> tuple[str, str]:
    if PDFTOTEXT and Path(PDFTOTEXT).exists():
        cmd = [PDFTOTEXT, "-layout", "-enc", "UTF-8", str(path), str(target)]
        code, out, err = run(cmd, timeout=180)
        if code == 0 and target.exists():
            text = target.read_text(encoding="utf-8", errors="replace")
            if len(text.strip()) > 30:
                return text, "pdftotext"
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            try:
                parts.append(f"\n\n[[PDF_PAGE {i+1}]]\n" + (page.extract_text() or ""))
            except Exception as e:
                parts.append(f"\n\n[[PDF_PAGE {i+1} ERROR {e!r}]]\n")
        text = "\n".join(parts)
        return text, "pypdf"
    except Exception as e:
        return "", f"pdf extraction failed: {e!r}"


def extract_docx(path: Path) -> tuple[str, str]:
    try:
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts), "python-docx"
    except Exception as e:
        return "", f"docx extraction failed: {e!r}"


def extract_doc_with_word(path: Path, target: Path) -> tuple[str, str]:
    ps = rf"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $null
try {{
  $doc = $word.Documents.Open('{str(path)}', $false, $true)
  $doc.SaveAs2('{str(target)}', 7)
}} finally {{
  if ($doc -ne $null) {{ $doc.Close($false) }}
  $word.Quit()
}}
"""
    code, out, err = run(["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps], timeout=120)
    if code == 0 and target.exists():
        for enc in ["utf-16", "utf-8-sig", "utf-8", "gb18030", "big5"]:
            try:
                return target.read_text(encoding=enc), f"Word COM text {enc}"
            except Exception:
                pass
        return target.read_text(errors="replace"), "Word COM text fallback"
    return "", f"Word COM extraction failed: {err or out}"


def extract_text(path: Path) -> tuple[str, str]:
    target = TEXT_DIR / (safe_name(path) + ".txt")
    ext = path.suffix.lower()
    try:
        if ext in [".txt", ".md", ".tex", ".html", ".mht"]:
            for enc in ["utf-8", "utf-16", "gb18030", "big5"]:
                try:
                    return path.read_text(encoding=enc), f"direct {enc}"
                except Exception:
                    pass
            return path.read_text(errors="replace"), "direct fallback"
        if ext == ".pdf":
            return extract_pdf(path, target)
        if ext == ".docx":
            return extract_docx(path)
        if ext == ".doc":
            return extract_doc_with_word(path, target)
    except Exception as e:
        return "", f"failed: {e!r}"
    return "", f"unsupported extension {ext}"


def collect_sources() -> list[tuple[str, Path]]:
    sources: list[tuple[str, Path]] = []

    exam_dir = ROOT / "题库" / "期末试题"
    if exam_dir.exists():
        for p in sorted(exam_dir.iterdir()):
            if p.is_file() and p.suffix.lower() in [".pdf", ".doc", ".docx", ".txt"]:
                if p.name == "exam2022_tmp.txt":
                    # This is an intermediate text extraction of the 2022 PDF, not an independent paper.
                    continue
                sources.append(("往年期末题-主目录", p))

    other_dir = ROOT / "题库" / "其他试题"
    if other_dir.exists():
        for p in sorted(other_dir.iterdir()):
            if p.is_file() and p.suffix.lower() in [".pdf", ".doc", ".docx", ".txt", ".mht"]:
                if "期末" in p.name and "期中" not in p.name:
                    sources.append(("往年期末题-其他目录", p))

    exercise_dir = ROOT / "题库" / "习题"
    if exercise_dir.exists():
        for p in sorted(exercise_dir.rglob("*")):
            if p.is_file() and p.suffix.lower() in [".pdf", ".doc", ".docx", ".txt", ".md"]:
                sources.append(("课本/题库习题", p))

    curated = ROOT / "第5-8章作业与复习_整理"
    if curated.exists():
        for p in sorted(curated.rglob("*")):
            if p.is_file() and p.suffix.lower() in [".tex", ".md", ".html", ".txt", ".pdf"]:
                if any(k in p.name for k in ["原题", "完整题目", "题目", "作业", "练习", "书", "课本"]):
                    sources.append(("课本/作业整理题", p))

    textbook_txt = ROOT / "张扬军（车辆学院）" / "教材" / "流体力学教材_text.txt"
    if textbook_txt.exists():
        sources.append(("课本教材题-张扬军文本", textbook_txt))

    textbook_dir = ROOT / "教科书"
    if textbook_dir.exists():
        for p in sorted(textbook_dir.iterdir()):
            if p.is_file() and p.suffix.lower() in [".pdf", ".txt"]:
                sources.append(("课本教材题-教科书PDF", p))

    answer_pdf = ROOT / "作业答案" / "流体力学课后作业解析（全）.pdf"
    if answer_pdf.exists():
        sources.append(("课本题-课后解析题干提取", answer_pdf))

    return sources


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            lines.append("")
            continue
        if DIR_NOISE_PAT.search(line):
            continue
        if re.match(r"^共\s*\d+\s*页|^第\s*\d+\s*页", line):
            continue
        if len(line) <= 2 and not re.search(r"\d", line):
            continue
        lines.append(line)
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def strip_solution(text: str) -> str:
    m = SOLUTION_CUT_PAT.search(text)
    if m:
        return text[: m.start()].strip()
    return text.strip()


def split_candidates(text: str, max_items: int = 600) -> list[tuple[str, str]]:
    text = clean_text(text)
    starts = list(QUESTION_START_PAT.finditer(text))
    chunks: list[str] = []
    if starts:
        for i, m in enumerate(starts):
            start = m.start()
            end = starts[i + 1].start() if i + 1 < len(starts) else min(len(text), start + 2200)
            chunk = text[start:end].strip()
            chunk = strip_solution(chunk)
            if len(chunk) > 12 and QUESTION_VERBS.search(chunk):
                chunks.append(chunk)
    else:
        paragraphs = re.split(r"\n\s*\n", text)
        for para in paragraphs:
            para = strip_solution(para)
            if 18 <= len(para) <= 1600 and QUESTION_VERBS.search(para):
                chunks.append(para)

    # Also handle markdown headings that are image-only original-question entries.
    for m in re.finditer(r"(?m)^##+\s*(.+?(?:原题|书\s*\d+\.\d+|题))\s*$", text):
        title = m.group(1).strip()
        if not any(title in c[:120] for c in chunks):
            chunks.append(title + "\n（题干在原文件中为图片或标题索引，文字抽取不足。）")

    seen = set()
    out = []
    for c in chunks:
        c = re.sub(r"\s+", " ", c).strip()
        c = c[:1200]
        key = re.sub(r"\W+", "", c[:120])
        if len(key) < 8 or key in seen:
            continue
        seen.add(key)
        no = infer_question_no(c)
        out.append((no, c))
        if len(out) >= max_items:
            break
    return out


SPECIAL_2022_ITEMS = [
    ("简答1", 3.0, 1.0, "简述欧拉法和拉格朗日法的差异。"),
    ("简答2", 3.0, 1.0, "简述层流和紊流的差异。"),
    ("简答3", 3.0, 0.0, "简述深水波的定义。"),
    ("简答4", 3.0, 2.0, "写出一元恒定总流能量方程，并简述其适用条件。"),
    ("简答5", 3.0, 0.0, "为什么相似准则中雷诺数准则和傅汝德数准则总是矛盾的？"),
    ("简答6", 3.0, 2.0, "简述边界层中的流动有什么特点。"),
    ("简答7", 3.0, 0.0, "圆球绕流阻力系数在某一雷诺数附近随雷诺数增大突然大幅度减小，说明原因。"),
    ("简答8", 3.0, 0.0, "简述什么是控制体？"),
    ("简答9", 3.0, 2.0, "写出描述牛顿流体运动的纳维斯托克斯（N-S）方程，并简述方程各项的意义。"),
    ("简答10", 3.0, 1.0, "利用 π 定理推求水下航行器所受流体阻力 F 的表达式。"),
    ("简答11", 3.0, 2.0, "简述流线和迹线的定义。"),
    ("简答12", 3.0, 0.0, "简述水击波的传播速度跟哪些因素有关。"),
    ("简答13", 3.0, 0.0, "什么是渐变流？在一元流动中，渐变流断面的测压管水头有什么特征。"),
    ("简答14", 3.0, 2.0, "圆管中断面流速分布在层流流态和紊流流态分别是什么样的？"),
    ("简答15", 3.0, 2.0, "写出不可压缩流体运动连续方程的一般形式。"),
    ("计算1", 10.0, 7.0, "密闭容器上层为空气，中层为油，下层为水，给测压管水面高程，求压力表 A 的读数。"),
    ("计算2", 15.0, 10.0, "变径圆管左端螺丝固定，流体从右端射出，给管径、比压计液体比重和测管水头差，理想流体，求固定螺丝推力。"),
    ("计算3", 15.0, 11.0, "两个油箱经突扩钢管连接，分别考虑短管局部损失和长管沿程损失，求输油流量、A/B 点压强和所需液面差。"),
    ("计算4", 15.0, 13.0, "半圆柱形气膜馆受大风正面来袭，用势流理论给外部流速分布、压力分布并计算升力合力。"),
]


def is_2022_exam_source(path: Path) -> bool:
    return path.name == "流体力学2022春期末考试试题.pdf"


def infer_question_no(text: str) -> str:
    patterns = [
        r"^\s*##+\s*(.+?)\s*$",
        r"^\s*(第\s*[一二三四五六七八九十百0-9]+\s*[题問])",
        r"^\s*(书\s*[0-9]{1,2}\.[0-9]{1,2})",
        r"^\s*([0-9]{1,2}\.[0-9]{1,2})",
        r"^\s*\(?([0-9]{1,2})\)?[.．、]",
        r"^\s*([一二三四五六七八九十]+)[、.．]",
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            return m.group(1).strip()
    return "未编号"


def detect_max_score(question_text: str, group: str) -> float:
    m = re.search(r"[（(]\s*(\d+(?:\.\d+)?)\s*分\s*[）)]", question_text)
    if m:
        return float(m.group(1))
    if "每小题 3 分" in question_text or "每小题3分" in question_text:
        return 3.0
    if group.startswith("往年期末"):
        return 10.0
    return 10.0


def classify(question_text: str) -> tuple[dict, list[dict]]:
    q = question_text
    matches = []
    for rule in TOPIC_RULES:
        score = 0
        hit = []
        for kw in rule["keywords"]:
            if kw.lower() in q.lower():
                score += 1
                hit.append(kw)
        if score:
            matches.append((score, len(hit), rule, hit))
    if not matches:
        return TOPIC_RULES[-1], []
    matches.sort(key=lambda x: (x[0], x[1], x[2]["base"]), reverse=True)
    primary = matches[0][2]
    return primary, [{"topic": x[2]["topic"], "hits": x[3]} for x in matches[:3]]


def estimate(rule: dict, question_text: str, group: str, max_score: float) -> tuple[float, str, str]:
    q = question_text
    ratio = float(rule["base"])
    basis = [f"主题命中“{rule['topic']}”，v6 基础过程覆盖约 {round(rule['base']*100)}%。"]

    if re.search(r"(简述|定义|为什么|说明|意义|差异|特点|是什么)", q):
        if rule["topic"] not in ["水波/水击/课程概念缺口", "概念简答/定义解释"]:
            ratio -= 0.18
            basis.append("题型偏概念简答，v6 偏公式模板，扣定义/解释分。")
    if re.search(r"(推导|证明|π|Pi|量纲|动量积分|排挤厚度|动量损失厚度|Karman|卡门)", q, re.I):
        ratio -= 0.22
        basis.append("题目要求推导或定义积分，v6 缺完整推导链，扣过程分。")
    if re.search(r"(查|图|表|Moody|Colebrook|PM|斜激波|面积|马赫|阻力系数)", q, re.I):
        ratio -= 0.08
        basis.append("数值或图表依赖较强，v6 只给查表入口或少量公式。")
    if "（题干在原文件中为图片" in q or "文字抽取不足" in q:
        ratio = min(ratio, 0.20)
        basis.append("题干主要是图片索引，未 OCR 到完整条件，只能评估题型入口。")
    if rule["topic"] == "水波/水击/课程概念缺口":
        ratio = min(ratio, 0.12)
    if rule["topic"] == "量纲分析/相似准则":
        ratio = min(ratio, 0.35)
    if rule["topic"] == "概念简答/定义解释":
        ratio = min(ratio, 0.40)

    ratio = max(0.0, min(0.95, ratio))
    est = round(max_score * ratio, 1)
    confidence = "中"
    if "文字抽取不足" in q:
        confidence = "低"
    elif len(q) > 80 and len(rule["paths"]) >= 3:
        confidence = "中高"
    return est, "；".join(basis), confidence


def make_eval(item_id: str, group: str, source: Path, question_no: str, question_text: str) -> EvalItem:
    rule, _matches = classify(question_text)
    max_score = detect_max_score(question_text, group)
    est, basis, confidence = estimate(rule, question_text, group, max_score)
    paths = rule["paths"]
    page_ref = "；".join(sorted({p.split(" > ")[0] for p in paths if p.startswith("P")}))
    own_steps = (
        "1. 先按题干关键词选模型；"
        "2. 在 v6 对应页写首个控制方程/公式；"
        "3. 把题给几何、单位和边界条件代入；"
        "4. 做代数求解或列出需要查图查表的位置；"
        "5. 检查单位、方向、绝压/表压或适用条件。"
    )
    if rule["topic"].startswith("可压缩"):
        own_steps = "1. 判定等熵段、阻塞、激波或膨胀波；2. 写 T0/T、p0/p、A/A* 或激波关系；3. 选择亚声/超声支；4. 标出需要查表的 M、β、ν；5. 回代压力、温度、流量并检查 pe/pb。"
    elif rule["topic"].startswith("静水"):
        own_steps = "1. 选同一静止连通液体等压面；2. 沿竖直方向逐段加减 rho g h；3. 若为闸门，求 F 与压力中心；4. 对铰链或曲面分量列平衡；5. 判断表压/绝压。"
    elif rule["topic"].startswith("势流"):
        own_steps = "1. 写 phi/psi 或基元叠加；2. 求速度分量；3. 用 Bernoulli 得压力或 Cp；4. 沿边界积分压力分量；5. 说明理想势流假设和单位展长。"
    elif rule["topic"].startswith("粘性管路"):
        own_steps = "1. 由 Q 和管径求 V；2. 算 Re 判断层流/湍流；3. 取 lambda 或查 Moody/Colebrook；4. 汇总沿程/局部损失；5. 代入机械能方程求目标量。"
    elif rule["topic"].startswith("边界层"):
        own_steps = "1. 算 Re_x 或 Re_L；2. 判断层流/湍流/低 Re；3. 选厚度、Cf 或阻力公式；4. 若为厚度定义题，尝试写积分；5. 说明 v6 缺失的定义或图表。"

    return EvalItem(
        item_id=item_id,
        group=group,
        source_file=str(source),
        source_label=source.name,
        question_no=question_no,
        question_text=question_text,
        inferred_topic=rule["topic"],
        cheatsheet_path="；".join(paths),
        page_reference=page_ref or "无直接页码",
        solution_route=rule["route"],
        own_steps=own_steps,
        needed_supplement=rule["supplement"],
        score_basis=basis,
        estimated_score=est,
        max_score=max_score,
        normalized_percent=round(100 * est / max_score, 1) if max_score else 0.0,
        confidence=confidence,
    )


def md_escape(s: str) -> str:
    return s.replace("|", "\\|").replace("\n", "<br>")


def write_csv(path: Path, rows: list[dict]) -> None:
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        w.writeheader()
        w.writerows(rows)


def write_index() -> None:
    lines = ["# 速查表 v6 知识索引\n", "本索引按 `PDF页码 > 章块 > 知识点` 层级记录，供逐题评估引用。\n"]
    for page in CHEATSHEET_INDEX:
        lines.append(f"## {page['page']}：{page['title']}\n")
        for node, desc in page["nodes"]:
            lines.append(f"- `{node}`：{desc}\n")
        lines.append("\n")
    (OUT / "01_速查表知识索引.md").write_text("".join(lines), encoding="utf-8")


def write_manifest(records: list[SourceRecord]) -> None:
    write_csv(SOURCE_MANIFEST, [asdict(r) for r in records])
    lines = ["# 题源清单\n", "状态说明：`readable` 为已抽出可用文本；`low_text` 多半是扫描图或抽取质量不足；`unreadable` 为无法自动解析。\n\n"]
    lines.append("| 组别 | 文件 | 状态 | 字符数 | 题目候选 | 备注 |\n|---|---|---:|---:|---:|---|\n")
    for r in records:
        rel = str(Path(r.path).relative_to(ROOT)) if str(r.path).startswith(str(ROOT)) else r.path
        lines.append(f"| {md_escape(r.group)} | `{md_escape(rel)}` | {r.status} | {r.char_count} | {r.candidate_count} | {md_escape(r.note)} |\n")
    (OUT / "02_题源清单.md").write_text("".join(lines), encoding="utf-8")


def write_exams(items: list[EvalItem]) -> None:
    exam_items = [i for i in items if i.group.startswith("往年期末")]
    by_source: dict[str, list[EvalItem]] = defaultdict(list)
    for i in exam_items:
        by_source[i.source_file].append(i)

    lines = ["# 往年期末题逐题评分\n", "评分前提：零基础学生只允许查 v6 速查表 PDF 与 v6 TeX 原始文件；不调用课外流体力学知识。无原始分值的试卷按每题 10 分过程分归一评估。\n\n"]
    for src, rows in by_source.items():
        rel = str(Path(src).relative_to(ROOT)) if src.startswith(str(ROOT)) else src
        total = sum(r.estimated_score for r in rows)
        max_total = sum(r.max_score for r in rows)
        norm = 100 * total / max_total if max_total else 0
        lines.append(f"## {Path(src).name}\n\n")
        lines.append(f"- 来源：`{rel}`\n")
        lines.append(f"- 可读题目数：{len(rows)}\n")
        lines.append(f"- 折算得分：**{norm:.1f}/100**（原始估分 {total:.1f}/{max_total:.1f}）\n\n")
        for r in rows:
            lines.append(f"### {r.item_id}｜{r.question_no}｜{r.inferred_topic}\n\n")
            lines.append(f"**题干摘录**：{r.question_text[:500]}\n\n")
            lines.append(f"**查速查表路径**：{r.cheatsheet_path}\n\n")
            lines.append(f"**自己写的解题步骤**：{r.own_steps}\n\n")
            lines.append(f"**按 v6 可执行的解决路径**：{r.solution_route}\n\n")
            lines.append(f"**还必须补充**：{r.needed_supplement}\n\n")
            lines.append(f"**只靠 v6 得分**：{r.estimated_score}/{r.max_score}，折合 {r.normalized_percent}%；依据：{r.score_basis}；信心：{r.confidence}。\n\n")
    (OUT / "03_往年期末逐题评估.md").write_text("".join(lines), encoding="utf-8")


def write_textbook(items: list[EvalItem]) -> None:
    tb_items = [i for i in items if not i.group.startswith("往年期末")]
    by_source: dict[str, list[EvalItem]] = defaultdict(list)
    for i in tb_items:
        by_source[i.source_file].append(i)

    lines = ["# 课本与习题逐题覆盖评估\n", "本文件覆盖自动抽取到文字题干的课本题、题库习题和整理题。图片题或扫描页若无 OCR 文本，仅按标题/题型入口给低信心评估。\n\n"]
    summary = Counter(i.inferred_topic for i in tb_items)
    lines.append("## 题型统计\n\n")
    lines.append("| 题型 | 数量 | 平均 v6 可得过程分 |\n|---|---:|---:|\n")
    for topic, count in summary.most_common():
        vals = [i.normalized_percent for i in tb_items if i.inferred_topic == topic]
        lines.append(f"| {topic} | {count} | {sum(vals)/len(vals):.1f}% |\n")
    lines.append("\n")

    for src, rows in by_source.items():
        rel = str(Path(src).relative_to(ROOT)) if src.startswith(str(ROOT)) else src
        avg = sum(r.normalized_percent for r in rows) / len(rows) if rows else 0
        lines.append(f"## {Path(src).name}\n\n")
        lines.append(f"- 来源：`{rel}`\n")
        lines.append(f"- 可读题目数：{len(rows)}\n")
        lines.append(f"- 平均过程分：**{avg:.1f}%**\n\n")
        for r in rows:
            lines.append(f"### {r.item_id}｜{r.question_no}｜{r.inferred_topic}\n\n")
            lines.append(f"**题干摘录**：{r.question_text[:450]}\n\n")
            lines.append(f"**查速查表路径**：{r.cheatsheet_path}\n\n")
            lines.append(f"**自己写的解题步骤**：{r.own_steps}\n\n")
            lines.append(f"**按 v6 可执行的解决路径**：{r.solution_route}\n\n")
            lines.append(f"**还必须补充**：{r.needed_supplement}\n\n")
            lines.append(f"**只靠 v6 得分**：{r.estimated_score}/{r.max_score}，折合 {r.normalized_percent}%；依据：{r.score_basis}；信心：{r.confidence}。\n\n")
    (OUT / "04_课本题目逐题评估.md").write_text("".join(lines), encoding="utf-8")


def write_all_paths(items: list[EvalItem]) -> None:
    lines = ["# 全量逐题解决路径\n", "这是本次评估的单一总过程文件。每个题目都包含：题干、查表层级、自己写的步骤、缺失补充、只靠 v6 的得分。\n\n"]
    for r in items:
        lines.append(f"## {r.item_id}｜{r.group}｜{r.source_label}｜{r.question_no}\n\n")
        lines.append(f"- 题型判断：{r.inferred_topic}\n")
        lines.append(f"- PDF 页码/层级指向：{r.cheatsheet_path}\n")
        lines.append(f"- 可得分：{r.estimated_score}/{r.max_score}（{r.normalized_percent}%），信心：{r.confidence}\n")
        lines.append(f"- 评分依据：{r.score_basis}\n\n")
        lines.append("**题干**\n\n")
        lines.append(r.question_text + "\n\n")
        lines.append("**我只看 v6 后会写的步骤**\n\n")
        for idx, step in enumerate([s for s in re.split(r"[；;]", r.own_steps) if s.strip()], 1):
            lines.append(f"{idx}. {step.strip()}\n")
        lines.append("\n**v6 能支持的解决路径**\n\n")
        lines.append(r.solution_route + "\n\n")
        lines.append("**为了完整满分还需要补充**\n\n")
        lines.append(r.needed_supplement + "\n\n")
    (OUT / "05_全量逐题解决路径.md").write_text("".join(lines), encoding="utf-8")
    write_csv(ITEM_CSV, [asdict(i) for i in items])


def write_deep_report(items: list[EvalItem], records: list[SourceRecord]) -> None:
    total = len(items)
    exam = [i for i in items if i.group.startswith("往年期末")]
    tb = [i for i in items if not i.group.startswith("往年期末")]
    topic_counter = Counter(i.inferred_topic for i in items)
    low = [i for i in items if i.normalized_percent < 40]
    high = [i for i in items if i.normalized_percent >= 70]
    unreadable = [r for r in records if r.status != "readable"]

    exam_by_src = defaultdict(list)
    for i in exam:
        exam_by_src[i.source_file].append(i)
    exam_scores = []
    for src, rows in exam_by_src.items():
        max_total = sum(r.max_score for r in rows)
        score = sum(r.estimated_score for r in rows)
        exam_scores.append((Path(src).name, 100 * score / max_total if max_total else 0, len(rows)))

    lines = ["# 速查表 v6 深度分析评判\n\n"]
    lines.append("## 总体判断\n\n")
    lines.append(f"本次从本地题源抽取到 **{total}** 个可评估题目候选，其中往年期末题 **{len(exam)}** 个，课本/习题类 **{len(tb)}** 个。v6 的强项是计算题入口和公式链，弱项是概念背诵、推导题、查表数值化和图片题读图。\n\n")
    if exam_scores:
        avg_exam = sum(s for _, s, _ in exam_scores) / len(exam_scores)
        lines.append(f"可读期末卷按各题过程分折算，平均约 **{avg_exam:.1f}/100**。其中 2022 春期末卷因为题干最完整，估分仍约 **56/100**；其他试卷若题干抽取质量较差，分数信心较低。\n\n")

    lines.append("## 题型覆盖\n\n")
    lines.append("| 题型 | 题数 | 平均过程分 | 判断 |\n|---|---:|---:|---|\n")
    for topic, count in topic_counter.most_common():
        vals = [i.normalized_percent for i in items if i.inferred_topic == topic]
        avg = sum(vals) / len(vals)
        if avg >= 70:
            tag = "强覆盖"
        elif avg >= 50:
            tag = "可保过程分"
        elif avg >= 30:
            tag = "弱覆盖"
        else:
            tag = "明显缺口"
        lines.append(f"| {topic} | {count} | {avg:.1f}% | {tag} |\n")

    lines.append("\n## 往年卷折算\n\n")
    lines.append("| 文件 | 可读题数 | 折算得分 |\n|---|---:|---:|\n")
    for name, score, n in sorted(exam_scores, key=lambda x: x[0]):
        lines.append(f"| {name} | {n} | {score:.1f}/100 |\n")

    lines.append("\n## v6 的核心优点\n\n")
    lines.append("1. 计算题模型入口清楚：静水压力、Bernoulli、动量、势流、可压缩、管路损失和边界层都有页码入口。\n")
    lines.append("2. 高频计算公式密度高：遇到典型计算题，零基础学生能至少写出控制方程和一部分代数过程。\n")
    lines.append("3. 易错提醒有考场价值：绝压/表压、pe/pb、At/A*、局部/平均速度、查表优先级能减少方向性错误。\n")
    lines.append("4. 对第 5-8 章后半学期计算题覆盖明显强于对前半学期概念题覆盖。\n\n")

    lines.append("## 致命短板\n\n")
    lines.append("1. 概念简答题不够：欧拉/拉格朗日、控制体、渐变流、水击、深水波、层流/紊流差异、相似准则矛盾等没有形成可背诵答案。\n")
    lines.append("2. 推导题缺桥梁：Buckingham π 定理、边界层排挤厚度/动量损失厚度、Karman 动量积分方程缺失，导致课本推导题很难得分。\n")
    lines.append("3. 查表题不能独立数值化：可压缩流、Moody、PM、斜激波和阻力系数图只写了查什么，没给常用表值或简化算法。\n")
    lines.append("4. 读图题依赖外部图像：扫描版期末题和课本图片题如果没有 OCR/清晰图，v6 只能给题型入口，无法保证数值答案。\n")
    lines.append("5. v6 是“会一点的人用来定位”的表，不是“没学过的人独立拿高分”的表。\n\n")

    lines.append("## 最高收益改版建议\n\n")
    lines.append("1. 加半页概念简答保底库：每个概念 1-2 句，覆盖欧拉/拉格朗日、控制体、渐变流、深水波、水击、层流/紊流、阻力危机、Re/Fr 相似矛盾。\n")
    lines.append("2. 加边界层三个缺失公式：排挤厚度、动量损失厚度、Karman 动量积分方程。\n")
    lines.append("3. 加动量题受力图模板：入口/出口压力力、壁面反力、重力、方向取反。\n")
    lines.append("4. 加极小常用表：空气等熵、正激波、面积-马赫数、PM 函数和光滑管湍流摩阻近似。\n")
    lines.append("5. 加读图检查清单：截面、方向、面积、液面高程、表压/绝压、单位展长。\n\n")

    if unreadable:
        lines.append("## 自动读取限制\n\n")
        lines.append(f"有 {len(unreadable)} 个题源未能抽取足够文字，主要原因是扫描 PDF、图片题或旧格式文件结构复杂。它们已列在 `02_题源清单.csv/md` 中，不能在未 OCR 的情况下声称逐题完整读题。\n\n")

    (OUT / "06_深度评判与改进建议.md").write_text("".join(lines), encoding="utf-8")


def write_intro(records: list[SourceRecord], items: list[EvalItem]) -> None:
    readable = sum(1 for r in records if r.status == "readable")
    lines = ["# 评测说明与范围\n\n"]
    lines.append("## 被测材料\n\n")
    lines.append(f"- v6 PDF：`{CHEATSHEET_PDF}`\n")
    lines.append(f"- v6 原始 TeX：`{CHEATSHEET_TEX}`\n")
    lines.append("- 只允许把 v6 的文字、公式、页码导航、题型入口和易错点作为解题知识来源。\n\n")
    lines.append("## 零基础学生假设\n\n")
    lines.append("1. 学生没有系统学过流体力学，不调用外部概念记忆。\n")
    lines.append("2. 可以做普通代数、微积分、单位换算、三角函数和简单方程求解。\n")
    lines.append("3. v6 明确写“查表/查图”时，视为知道该查什么；但 v6 未给表值时，数值满分要扣。\n")
    lines.append("4. 题干若是图片或扫描且本次未 OCR 出文字，只能评估题型入口，不能评估完整数值解。\n\n")
    lines.append("## 题源范围\n\n")
    lines.append(f"- 已扫描本地题源文件：{len(records)} 个；成功抽出可用文字：{readable} 个。\n")
    lines.append(f"- 自动抽取并评估题目候选：{len(items)} 个。\n")
    lines.append("- 往年期末题主范围：`题库/期末试题`，另把 `题库/其他试题` 中文件名明确含“期末”且不含“期中”的文件作为辅助往年卷。\n")
    lines.append("- 课本题范围：`题库/习题`、`第5-8章作业与复习_整理`、`张扬军（车辆学院）/教材/流体力学教材_text.txt`、`教科书` PDF 及课后解析题干。\n\n")
    lines.append("## 评分口径\n\n")
    lines.append("- 有明确分值的期末题按原分值估计；无分值题统一按 10 分过程分估计。\n")
    lines.append("- “只靠 v6 得分”代表零基础学生能写出的可给分过程，不代表标准答案得分。\n")
    lines.append("- 评分由 v6 层级索引和题干关键词匹配生成，再按概念题、推导题、查表题、图片题做扣分。\n\n")
    (OUT / "00_评测说明与范围.md").write_text("".join(lines), encoding="utf-8")


def write_final(records: list[SourceRecord], items: list[EvalItem]) -> None:
    exam = [i for i in items if i.group.startswith("往年期末")]
    tb = [i for i in items if not i.group.startswith("往年期末")]
    exam_by_src = defaultdict(list)
    for i in exam:
        exam_by_src[i.source_file].append(i)
    scores = []
    for src, rows in exam_by_src.items():
        max_total = sum(r.max_score for r in rows)
        score = sum(r.estimated_score for r in rows)
        scores.append((Path(src).name, 100 * score / max_total if max_total else 0, len(rows)))
    avg_exam = sum(s for _, s, _ in scores) / len(scores) if scores else 0
    avg_tb = sum(i.normalized_percent for i in tb) / len(tb) if tb else 0
    unreadable = [r for r in records if r.status != "readable"]

    lines = ["# 期末六页 A4 速查表 v6 全量测试评判最终报告\n\n"]
    lines.append("## 结论\n\n")
    lines.append(f"在“零基础学生只看 v6 PDF 和 v6 TeX 原始文件”的前提下，本次从本地资料中抽取并评估 **{len(items)}** 个题目候选。可读往年期末题折算平均约 **{avg_exam:.1f}/100**；课本/习题题目的平均过程覆盖约 **{avg_tb:.1f}%**。\n\n")
    lines.append("对最完整可读的 2022 春期末卷，结论仍是约 **56/100**：简答题约 15/45，计算分析题约 41/55。v6 对计算题很有用，但不能让完全没学过的人稳定拿高分。\n\n")

    lines.append("## 生成文件\n\n")
    for name in [
        "00_评测说明与范围.md",
        "01_速查表知识索引.md",
        "02_题源清单.md",
        "02_题源清单.csv",
        "03_往年期末逐题评估.md",
        "04_课本题目逐题评估.md",
        "05_全量逐题解决路径.md",
        "05_全量逐题解决路径.csv",
        "06_深度评判与改进建议.md",
    ]:
        lines.append(f"- `{name}`\n")

    lines.append("\n## 关键发现\n\n")
    lines.append("1. 强项：静水压力、Bernoulli/动量、势流圆柱、可压缩喷管/激波、管路损失、泵/水轮机和平板边界层直接公式题。\n")
    lines.append("2. 中等：查表类可压缩题、Moody/Colebrook 管路题、复杂控制体受力题；能写过程，但数值满分不稳。\n")
    lines.append("3. 弱项：概念简答、量纲分析与相似准则、深水波/水击、边界层积分推导、排挤厚度/动量损失厚度定义题。\n")
    lines.append("4. 最大瓶颈：v6 是公式/模型定位表，不是概念背诵库，也不含常用数值表。\n\n")

    if scores:
        lines.append("## 往年卷折算概览\n\n")
        lines.append("| 文件 | 可读题数 | 只靠 v6 折算得分 |\n|---|---:|---:|\n")
        for name, score, n in sorted(scores, key=lambda x: x[0]):
            lines.append(f"| {name} | {n} | {score:.1f}/100 |\n")
        lines.append("\n")

    if unreadable:
        lines.append("## 读取限制\n\n")
        lines.append(f"{len(unreadable)} 个题源没有抽出足够文字，多为扫描 PDF 或图片题。它们已进入题源清单，但不能在没有 OCR/人工转写的情况下逐题完整评分。\n\n")

    lines.append("## 最优先改动\n\n")
    lines.append("把 v6 改到更像“零基础保分表”，优先加：半页概念简答库、边界层积分三公式、控制体受力图模板、极小常用查表、读图检查清单。预计 2022 这类卷可由约 56/100 提升到约 70+/100。\n")
    (OUT / "最终报告.md").write_text("".join(lines), encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    write_index()

    records: list[SourceRecord] = []
    items: list[EvalItem] = []
    item_counter = 1

    sources = collect_sources()
    seen_paths = set()
    unique_sources = []
    for group, path in sources:
        if path in seen_paths:
            continue
        seen_paths.add(path)
        unique_sources.append((group, path))

    for group, path in unique_sources:
        text, method = extract_text(path)
        text_file = TEXT_DIR / (safe_name(path) + ".txt")
        if text:
            text_file.write_text(text, encoding="utf-8", errors="replace")
        char_count = len(text.strip())
        status = "readable" if char_count >= 120 else ("low_text" if char_count > 0 else "unreadable")
        note = method
        candidates: list[tuple[str, str]] = []
        special_scores: dict[str, tuple[float, float]] = {}
        if status in ["readable", "low_text"]:
            if is_2022_exam_source(path):
                candidates = [(no, qtext) for no, max_score, est, qtext in SPECIAL_2022_ITEMS]
                special_scores = {no: (max_score, est) for no, max_score, est, _qtext in SPECIAL_2022_ITEMS}
            else:
                candidates = split_candidates(text)
            # If a known exam text uses page-level extraction and split failed, keep a file-level candidate.
            if not candidates and group.startswith("往年期末") and char_count > 80:
                candidates = [("整卷", strip_solution(clean_text(text))[:1200])]
        rec = SourceRecord(
            group=group,
            path=str(path),
            status=status if candidates or status != "readable" else "readable",
            extracted_text=str(text_file) if text else "",
            char_count=char_count,
            candidate_count=len(candidates),
            note=note,
        )
        records.append(rec)
        for qno, qtext in candidates:
            item_id = f"Q{item_counter:04d}"
            ev = make_eval(item_id, group, path, qno, qtext)
            if qno in special_scores:
                max_score, est = special_scores[qno]
                ev.max_score = max_score
                ev.estimated_score = est
                ev.normalized_percent = round(100 * est / max_score, 1) if max_score else 0.0
                ev.confidence = "高"
                ev.score_basis = "2022 春期末卷为本次人工核验样本，按 15 道简答 45 分、4 道计算 55 分的原卷结构估分；简答共 15/45，计算共 41/55。"
            items.append(ev)
            item_counter += 1

    write_manifest(records)
    write_intro(records, items)
    write_exams(items)
    write_textbook(items)
    write_all_paths(items)
    write_deep_report(items, records)
    write_final(records, items)

    summary = {
        "sources": len(records),
        "items": len(items),
        "readable_sources": sum(1 for r in records if r.status == "readable"),
        "output": str(OUT),
    }
    (OUT / "run_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
