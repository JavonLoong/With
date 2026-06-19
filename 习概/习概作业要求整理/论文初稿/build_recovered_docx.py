# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


BASE_DIR = Path(__file__).resolve().parent
TEX_PATH = BASE_DIR / "习概期中作业_恢复_2128版.tex"
DOCX_PATH = BASE_DIR / "习概期中作业_恢复_2128版_圈码脚注版.docx"
TMP_DOCX_PATH = BASE_DIR / "习概期中作业_恢复_2128版_圈码脚注版.tmp.docx"

TITLE = "从《摆脱贫困》看以人民为中心的发展思想"
SUBTITLE = "——兼论精准扶贫的基层治理逻辑"
META = "姓名：纪文龙    学号：2024012842    班级：行健烽火4"

FOOTNOTES: dict[int, str] = {}
CIRCLED_NUMBERS = {
    1: "①",
    2: "②",
    3: "③",
    4: "④",
    5: "⑤",
    6: "⑥",
    7: "⑦",
    8: "⑧",
    9: "⑨",
    10: "⑩",
}


def set_run_font(run, font_name="宋体", size=None, bold=None):
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def apply_body_format(paragraph, first_line=True, after=0, line=1.35):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_footnote_reference(paragraph, note_id):
    mark = CIRCLED_NUMBERS.get(note_id, str(note_id))

    run_element = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "FootnoteReference")
    rpr.append(style)
    vert = OxmlElement("w:vertAlign")
    vert.set(qn("w:val"), "superscript")
    rpr.append(vert)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    rpr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "16")
    rpr.append(sz_cs)
    run_element.append(rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(note_id))
    ref.set(qn("w:customMarkFollows"), "1")
    run_element.append(ref)
    paragraph._p.append(run_element)

    mark_run = paragraph.add_run(mark)
    set_run_font(mark_run, "宋体", 8)
    mark_run.font.superscript = True


def latex_clean(text: str) -> str:
    text = text.replace("\\%", "%")
    text = text.replace("\\&", "&")
    text = text.replace("\\_", "_")
    text = text.replace("\\par", "")
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+(\[[^\]]*\])?(\{[^{}]*\})?", "", text)
    text = text.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", text).strip()


def parse_parts(text: str):
    parts = []
    pos = 0
    note_id = len(FOOTNOTES) + 1
    pattern = re.compile(r"\\footnote\{([^{}]*)\}")
    for match in pattern.finditer(text):
        before = latex_clean(text[pos : match.start()])
        if before:
            parts.append(before)
        FOOTNOTES[note_id] = latex_clean(match.group(1))
        parts.append(note_id)
        note_id += 1
        pos = match.end()
    tail = latex_clean(text[pos:])
    if tail:
        parts.append(tail)
    return parts


def add_mixed_paragraph(doc, parts, *, first_line=True, after=0):
    p = doc.add_paragraph()
    apply_body_format(p, first_line=first_line, after=after)
    for part in parts:
        if isinstance(part, int):
            add_footnote_reference(p, part)
        else:
            run = p.add_run(part)
            set_run_font(run, "宋体", 12)
    return p


def extract_body(tex: str):
    abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", tex, re.S).group(1).strip()
    keyword_match = re.search(r"\\noindent\s*\\textbf\{关键词：\}(.*)", tex)
    keywords = keyword_match.group(1).strip() if keyword_match else ""

    start = tex.index("\\section*")
    end = tex.index("\\end{document}")
    body = tex[start:end]
    pieces = re.split(r"\\section\*\{([^{}]+)\}", body)
    sections = []
    for idx in range(1, len(pieces), 2):
        heading = latex_clean(pieces[idx])
        content = pieces[idx + 1]
        content = re.sub(r"\\vspace\{[^{}]*\}", "", content)
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", content) if p.strip()]
        sections.append((heading, paragraphs))
    return abstract, keywords, sections


def build_docx_base(path: Path):
    tex = TEX_PATH.read_text(encoding="utf-8")
    abstract, keywords, sections = extract_body(tex)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run(TITLE)
    set_run_font(r, "黑体", 22, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run(SUBTITLE)
    set_run_font(r, "黑体", 16, True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    r = meta.add_run(META)
    set_run_font(r, "宋体", 10.5)

    abstract_p = doc.add_paragraph()
    apply_body_format(abstract_p, first_line=False, after=6)
    label = abstract_p.add_run("摘 要：")
    set_run_font(label, "黑体", 12, True)
    for idx, para in enumerate([p.strip() for p in abstract.splitlines() if p.strip()]):
        if idx:
            abstract_p.add_run("\n")
        r = abstract_p.add_run(latex_clean(para))
        set_run_font(r, "宋体", 12)

    keyword_p = doc.add_paragraph()
    apply_body_format(keyword_p, first_line=False, after=12)
    label = keyword_p.add_run("关键词：")
    set_run_font(label, "黑体", 12, True)
    r = keyword_p.add_run(latex_clean(keywords))
    set_run_font(r, "宋体", 12)

    for heading, paragraphs in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(heading)
        set_run_font(r, "黑体", 13, True)
        for paragraph in paragraphs:
            add_mixed_paragraph(doc, parse_parts(paragraph), after=0)

    doc.save(path)


def make_footnotes_xml():
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xml_ns = "http://www.w3.org/XML/1998/namespace"
    footnotes = etree.Element(f"{{{w}}}footnotes", nsmap={"w": w})

    sep = etree.SubElement(footnotes, f"{{{w}}}footnote")
    sep.set(f"{{{w}}}type", "separator")
    sep.set(f"{{{w}}}id", "-1")
    p = etree.SubElement(sep, f"{{{w}}}p")
    r = etree.SubElement(p, f"{{{w}}}r")
    etree.SubElement(r, f"{{{w}}}separator")

    cont = etree.SubElement(footnotes, f"{{{w}}}footnote")
    cont.set(f"{{{w}}}type", "continuationSeparator")
    cont.set(f"{{{w}}}id", "0")
    p = etree.SubElement(cont, f"{{{w}}}p")
    r = etree.SubElement(p, f"{{{w}}}r")
    etree.SubElement(r, f"{{{w}}}continuationSeparator")

    for note_id, text in FOOTNOTES.items():
        mark = CIRCLED_NUMBERS.get(note_id, str(note_id))
        fn = etree.SubElement(footnotes, f"{{{w}}}footnote")
        fn.set(f"{{{w}}}id", str(note_id))
        p = etree.SubElement(fn, f"{{{w}}}p")
        ppr = etree.SubElement(p, f"{{{w}}}pPr")
        pstyle = etree.SubElement(ppr, f"{{{w}}}pStyle")
        pstyle.set(f"{{{w}}}val", "FootnoteText")
        r = etree.SubElement(p, f"{{{w}}}r")
        rpr = etree.SubElement(r, f"{{{w}}}rPr")
        rstyle = etree.SubElement(rpr, f"{{{w}}}rStyle")
        rstyle.set(f"{{{w}}}val", "FootnoteReference")
        vert = etree.SubElement(rpr, f"{{{w}}}vertAlign")
        vert.set(f"{{{w}}}val", "superscript")
        sz = etree.SubElement(rpr, f"{{{w}}}sz")
        sz.set(f"{{{w}}}val", "16")
        sz_cs = etree.SubElement(rpr, f"{{{w}}}szCs")
        sz_cs.set(f"{{{w}}}val", "16")
        tmark = etree.SubElement(r, f"{{{w}}}t")
        tmark.text = mark
        text_run = etree.SubElement(p, f"{{{w}}}r")
        text_rpr = etree.SubElement(text_run, f"{{{w}}}rPr")
        fonts = etree.SubElement(text_rpr, f"{{{w}}}rFonts")
        fonts.set(f"{{{w}}}eastAsia", "宋体")
        fonts.set(f"{{{w}}}ascii", "Times New Roman")
        fonts.set(f"{{{w}}}hAnsi", "Times New Roman")
        sz = etree.SubElement(text_rpr, f"{{{w}}}sz")
        sz.set(f"{{{w}}}val", "18")
        sz_cs = etree.SubElement(text_rpr, f"{{{w}}}szCs")
        sz_cs.set(f"{{{w}}}val", "18")
        t = etree.SubElement(text_run, f"{{{w}}}t")
        t.set(f"{{{xml_ns}}}space", "preserve")
        t.text = " " + text

    return etree.tostring(footnotes, xml_declaration=True, encoding="UTF-8", standalone="yes")


def patch_docx_footnotes(src: Path, dst: Path):
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    foot_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp) / "base.docx"
        shutil.copyfile(src, tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
            content_types = etree.fromstring(zin.read("[Content_Types].xml"))
            existing = content_types.xpath(
                "ct:Override[@PartName='/word/footnotes.xml']",
                namespaces={"ct": ct_ns},
            )
            if not existing:
                override = etree.SubElement(content_types, f"{{{ct_ns}}}Override")
                override.set("PartName", "/word/footnotes.xml")
                override.set(
                    "ContentType",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
                )

            rels = etree.fromstring(zin.read("word/_rels/document.xml.rels"))
            if not rels.xpath("rel:Relationship[@Type=$rtype]", namespaces={"rel": rel_ns}, rtype=foot_rel_type):
                ids = []
                for rel in rels.xpath("rel:Relationship", namespaces={"rel": rel_ns}):
                    rid = rel.get("Id", "")
                    match = re.match(r"rId(\d+)$", rid)
                    if match:
                        ids.append(int(match.group(1)))
                rel = etree.SubElement(rels, f"{{{rel_ns}}}Relationship")
                rel.set("Id", f"rId{max(ids or [0]) + 1}")
                rel.set("Type", foot_rel_type)
                rel.set("Target", "footnotes.xml")

            for item in zin.infolist():
                if item.filename in {
                    "[Content_Types].xml",
                    "word/_rels/document.xml.rels",
                    "word/footnotes.xml",
                }:
                    continue
                zout.writestr(item, zin.read(item.filename))

            zout.writestr(
                "[Content_Types].xml",
                etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
            zout.writestr(
                "word/_rels/document.xml.rels",
                etree.tostring(rels, xml_declaration=True, encoding="UTF-8", standalone=True),
            )
            zout.writestr("word/footnotes.xml", make_footnotes_xml())


def main():
    build_docx_base(TMP_DOCX_PATH)
    patch_docx_footnotes(TMP_DOCX_PATH, DOCX_PATH)
    TMP_DOCX_PATH.unlink(missing_ok=True)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
